"""
Dialog Windows - File Converter Pro

Collection of modal and non-modal dialog windows for user interaction.

Main Dialogs:
    - SettingsDialog: Application preferences and configuration.
    - TermsAndPrivacyDialog: Legal acceptance (CGU/Privacy Policy).  →  terms_dialog.py
    - ModernSplashScreen: Animated loading screen with progress steps.
    - CompressionDialog: Archive settings (ZIP/RAR/TAR, split, encrypt).
    - SplitDialog: PDF splitting options (pages, ranges).
    - PasswordDialog: Secure password input for PDF protection.
    - PreviewDialog: Quick file preview (PDF, Images, Word).
    - ConversionOptionsDialog: Generic conversion selection menu.
    - WordToPdfOptionsDialog: Specific options for Word→PDF conversion(from word_to_pdf_dialog.py).

Features:
    - Bilingual support via TranslationManager
    - Theme-aware CSS styling
    - Robust resource path handling (Dev + PyInstaller)

Author: Hyacinthe
Version: 1.0
"""

import sys
import os
import fitz
from pathlib import Path
from PySide6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QPushButton, 
                               QLabel, QFileDialog, QMessageBox,
                               QProgressBar, QComboBox, QFrame,
                               QGroupBox, QScrollArea, QLineEdit, QDialog,
                               QDialogButtonBox, QFormLayout, QSpinBox,
                               QTabWidget, QRadioButton, QButtonGroup,
                               QGridLayout, QTableWidget,
                               QTableWidgetItem, QHeaderView)
from PySide6.QtCore import (Qt, QPropertyAnimation, QEasingCurve, QTimer, QSize,
                            QSequentialAnimationGroup, QCoreApplication, Signal)
from PySide6.QtGui import (QIcon, QPixmap)
import sys as _sys, os as _os
_PKG_DIR  = _os.path.dirname(_os.path.abspath(__file__))
_ROOT_DIR = _os.path.dirname(_PKG_DIR)
if _ROOT_DIR not in _sys.path:
    _sys.path.insert(0, _ROOT_DIR)

from datetime import datetime

# Local imports
from widgets import AnimatedCheckBox
from .terms_dialog import TermsAndPrivacyDialog
from translations import TranslationManager

def _make_tm(language):
    """Helper: create a TranslationManager set to *language*."""
    tm = TranslationManager()
    tm.set_language(language)
    return tm

class PdfToWordDialog(QDialog):
    def __init__(self, parent=None, language="fr", current_mode="with_images", has_images=False):
        super().__init__(parent)
        self.language = language
        self._tm = _make_tm(language)
        self.current_mode = current_mode
        self.has_images = has_images
        self.setWindowTitle(self.translate_text("Options de conversion PDF vers Word"))
        self.setModal(True)
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)

        mode_group = QGroupBox(self.translate_text("Mode de conversion"))
        mode_layout = QVBoxLayout(mode_group)

        self.mode_group = QButtonGroup(self)

        self.with_images_radio = QRadioButton(self.translate_text("Conserver les images et la mise en page (recommandé)"))
        self.text_only_radio = QRadioButton(self.translate_text("Texte brut uniquement (plus rapide)"))
        self.text_with_image_text_radio = QRadioButton(self.translate_text("Texte complet (texte + texte des images)"))

        self.mode_group.addButton(self.with_images_radio, 1)
        self.mode_group.addButton(self.text_only_radio, 2)
        self.mode_group.addButton(self.text_with_image_text_radio, 3)

        if self.current_mode == "with_images":
            self.with_images_radio.setChecked(True)
        elif self.current_mode == "text_only":
            self.text_only_radio.setChecked(True)
        else:
            self.text_with_image_text_radio.setChecked(True)

        mode_layout.addWidget(self.with_images_radio)
        mode_layout.addWidget(self.text_only_radio)

        if self.has_images:
            mode_layout.addWidget(self.text_with_image_text_radio)
            self.image_info_label = QLabel(self.translate_text("ℹ️ Ce PDF contient des images. L'option 'Texte complet' extraira le texte des images."))
            self.image_info_label.setStyleSheet("color: #007acc; font-size: 11px; margin-top: 10px;")
            self.image_info_label.setWordWrap(True)
            mode_layout.addWidget(self.image_info_label)
        else:
            self.image_info_label = QLabel(self.translate_text("ℹ️ Ce PDF ne contient pas d'images détectées."))
            self.image_info_label.setStyleSheet("color: #28a745; font-size: 11px; margin-top: 10px;")
            self.image_info_label.setWordWrap(True)
            mode_layout.addWidget(self.image_info_label)

        layout.addWidget(mode_group)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        ok_button = buttons.button(QDialogButtonBox.Ok)
        if ok_button:
            ok_button.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
                QPushButton:pressed {
                    background-color: #1e7e34;
                }
            """)
        cancel_button = buttons.button(QDialogButtonBox.Cancel)
        if cancel_button:
            cancel_button.setStyleSheet("""
                QPushButton {
                    background-color: #B55454;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #A04040;
                }
                QPushButton:pressed {
                    background-color: #8B3030;
                }
            """)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout.addWidget(buttons)

    def translate_text(self, text):
        return self._tm.translate_text(text)

    def get_conversion_mode(self):
        if self.with_images_radio.isChecked():
            return "with_images"
        elif self.text_only_radio.isChecked():
            return "text_only"
        else:
            return "text_with_image_text"

class ModernSplashScreen(QWidget):
    def __init__(self, config, parent=None):
        super().__init__(parent)
        self.config = config
        self.dark_mode = config.get("dark_mode", False)
        self.current_language = config.get("language", "fr")
        self._tm = _make_tm(self.current_language)
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setFixedSize(500, 400)
        self.setWindowIcon(QIcon(self.get_icon_path()))
        self.setup_ui()
        self.setWindowTitle(self.translate_text("File Converter Pro - convertisseur de fichiers professionnels"))

    def get_icon_path(self):
        """Retrieve icon path (icon.ico) robustly (dev + PyInstaller)"""
        icon_name = "icon.ico"

        if getattr(sys, 'frozen', False):
            base_path = sys._MEIPASS
            path = os.path.join(base_path, icon_name)
            if os.path.exists(path):
                return path

        path = os.path.join(_ROOT_DIR, icon_name)
        if os.path.exists(path):
            return path

        path = os.path.join(os.getcwd(), icon_name)
        if os.path.exists(path):
            return path

        return icon_name

    def setup_ui(self):
        layout = QVBoxLayout()
        layout.setAlignment(Qt.AlignCenter)
        layout.setContentsMargins(0, 0, 0, 0)

        self.container = QFrame()
        self.container.setObjectName("container")
        container_layout = QVBoxLayout(self.container)
        container_layout.setAlignment(Qt.AlignCenter)
        container_layout.setSpacing(20)

        self.logo_container = QWidget()
        self.logo_container.setFixedSize(120, 120)
        logo_layout = QVBoxLayout(self.logo_container)
        logo_layout.setAlignment(Qt.AlignCenter)

        self.logo_label = QLabel("📦")
        self.logo_label.setAlignment(Qt.AlignCenter)
        self.logo_label.setStyleSheet("""
            QLabel {
                font-size: 60px;
                background: transparent;
            }
        """)
        logo_layout.addWidget(self.logo_label)

        self.title_label = QLabel("FILE CONVERTER PRO")
        self.title_label.setAlignment(Qt.AlignCenter)
        self.title_label.setStyleSheet("""
            QLabel {
                font-size: 32px;
                font-weight: bold;
                font-family: 'Segoe UI', Arial, sans-serif;
            }
        """)

        subtitle_text = self.get_translated_text("Initialisation de l'application...")
        self.subtitle_label = QLabel(subtitle_text)
        self.subtitle_label.setAlignment(Qt.AlignCenter)
        self.subtitle_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                margin: 5px;
            }
        """)

        self.progress_container = QWidget()
        self.progress_container.setFixedWidth(300)
        progress_layout = QVBoxLayout(self.progress_container)

        self.progress_bar = QProgressBar()
        self.progress_bar.setMaximum(100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setFixedHeight(8)

        self.progress_text = QLabel("0%")
        self.progress_text.setAlignment(Qt.AlignCenter)
        self.progress_text.setStyleSheet("font-size: 12px; margin-top: 5px;")

        progress_layout.addWidget(self.progress_bar)
        progress_layout.addWidget(self.progress_text)

        container_layout.addWidget(self.logo_container)
        container_layout.addWidget(self.title_label)
        container_layout.addWidget(self.subtitle_label)
        container_layout.addWidget(self.progress_container)

        layout.addWidget(self.container)
        self.setLayout(layout)

        self.apply_styles()
        self.setup_animations()

    def translate_text(self, text):
        return self._tm.translate_text(text)

    def get_translated_text(self, text):
        return self._tm.translate_text(text)

    def apply_styles(self):
        if self.dark_mode:
            self.setStyleSheet("""
                #container {
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                        stop:0 #2d333b, stop:1 #1a1d23);
                    border-radius: 20px;
                    border: 1px solid #495057;
                }
                QLabel {
                    color: #e9ecef;
                }
                QProgressBar {
                    background-color: #495057;
                    border: none;
                    border-radius: 4px;
                }
                QProgressBar::chunk {
                    background-color: #4dabf7;
                    border-radius: 4px;
                }
            """)
            self.title_label.setStyleSheet("color: #e9ecef; font-size: 32px; font-weight: bold;")
            self.subtitle_label.setStyleSheet("color: #adb5bd; font-size: 14px;")
            self.progress_text.setStyleSheet("color: #adb5bd; font-size: 12px;")
        else:
            self.setStyleSheet("""
                #container {
                    background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                        stop:0 #ffffff, stop:1 #f8f9fa);
                    border-radius: 20px;
                    border: 1px solid #e0e0e0;
                }
                QLabel {
                    color: #2c3e50;
                }
                QProgressBar {
                    background-color: #ecf0f1;
                    border: none;
                    border-radius: 4px;
                }
                QProgressBar::chunk {
                    background-color: #3498db;
                    border-radius: 4px;
                }
            """)
            self.title_label.setStyleSheet("color: #2c3e50; font-size: 32px; font-weight: bold;")
            self.subtitle_label.setStyleSheet("color: #7f8c8d; font-size: 14px;")
            self.progress_text.setStyleSheet("color: #7f8c8d; font-size: 12px;")

    def setup_animations(self):
        self.progress_animation = QPropertyAnimation(self.progress_bar, b"value")
        self.progress_animation.setDuration(2500)
        self.progress_animation.setStartValue(0)
        self.progress_animation.setEndValue(100)
        self.progress_animation.setEasingCurve(QEasingCurve.OutCubic)

        self.progress_animation.valueChanged.connect(self.update_progress_text)

    def update_progress_text(self, value):
        self.progress_text.setText(f"{int(value)}%")

        if value < 25:
            current_text = "Initialisation de l'application..."
        elif value < 50:
            current_text = "Chargement des modules..."
        elif value < 75:
            current_text = "Préparation de l'interface..."
        elif value < 95:
            current_text = "Presque terminé..."
        else:
            current_text = "Prêt!"

        self.subtitle_label.setText(self.get_translated_text(current_text))

    def start_animation(self):
        logo_animation = QPropertyAnimation(self.logo_label, b"geometry")
        logo_animation.setDuration(1000)
        logo_animation.setKeyValueAt(0, self.logo_label.geometry())
        logo_animation.setKeyValueAt(0.3, self.logo_label.geometry().adjusted(0, -15, 0, -15))
        logo_animation.setKeyValueAt(1, self.logo_label.geometry())
        logo_animation.setEasingCurve(QEasingCurve.OutBounce)

        self.animation_group = QSequentialAnimationGroup()
        self.animation_group.addAnimation(logo_animation)
        self.animation_group.addAnimation(self.progress_animation)
        self.animation_group.finished.connect(self.close)
        self.animation_group.finished.connect(self.deleteLater)
        self.animation_group.start()

    def fade_out(self, duration=500):
        """Launch a smooth fade-out animation and destroy the window"""
        self.fade_animation = QPropertyAnimation(self, b"windowOpacity")
        self.fade_animation.setDuration(duration)
        self.fade_animation.setStartValue(1.0)
        self.fade_animation.setEndValue(0.0)
        self.fade_animation.setEasingCurve(QEasingCurve.OutCubic)
        self.fade_animation.finished.connect(self.close)
        self.fade_animation.finished.connect(self.deleteLater)
        self.fade_animation.start()

class PreviewDialog(QDialog):
    def __init__(self, file_path, parent=None, language="fr"):
        super().__init__(parent)
        self.file_path = file_path
        self.language = language
        self._tm = _make_tm(language)
        self._media_player = None  # QMediaPlayer ref kept alive
        self.setWindowTitle(self.translate_text("Aperçu -") + f" {Path(file_path).name}")
        self.setModal(False)
        self.setMinimumSize(700, 600)
        self.resize(800, 680)
        self.setup_ui()
        self.load_preview()

    def translate_text(self, text):
        return self._tm.translate_text(text)

    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        self.header_label = QLabel(f"<b>{Path(self.file_path).name}</b>")
        self.header_label.setStyleSheet("font-size: 14px; padding: 4px 8px;")
        layout.addWidget(self.header_label)

        # Central content area — filled by each preview_* method
        self.content_area = QWidget()
        self.content_layout = QVBoxLayout(self.content_area)
        self.content_layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.content_area, 1)

        # Fallback label used by simple previews
        self.preview_label = QLabel()
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setWordWrap(True)
        self.content_layout.addWidget(self.preview_label)

        self.close_btn = QPushButton(self.translate_text("Fermer l'aperçu"))
        self.close_btn.setFixedHeight(36)
        self.close_btn.clicked.connect(self.close)
        layout.addWidget(self.close_btn)

    # helpers
    def _clear_content(self):
        """Remove all widgets from content_layout (except preview_label)."""
        while self.content_layout.count():
            item = self.content_layout.takeAt(0)
            w = item.widget()
            if w and w is not self.preview_label:
                w.deleteLater()
        self.preview_label.clear()
        self.preview_label.setParent(None)  # detach, will be re-added if needed

    def _show_error(self, msg):
        self.preview_label.setParent(self.content_area)
        self.content_layout.addWidget(self.preview_label)
        self.preview_label.setText(f"⚠️ {msg}")

    # router
    def load_preview(self):
        ext = Path(self.file_path).suffix.lower()
        try:
            if ext == '.pdf':
                self.preview_pdf()
            elif ext in ('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp', '.gif'):
                self.preview_image()
            elif ext in ('.docx', '.doc'):
                self.preview_word()
            elif ext in ('.txt', '.log', '.md', '.py', '.js', '.ts', '.css', '.xml', '.yaml', '.yml', '.ini', '.cfg', '.bat', '.sh'):
                self.preview_text(syntax=ext)
            elif ext == '.rtf':
                self.preview_rtf()
            elif ext == '.csv':
                self.preview_csv()
            elif ext == '.json':
                self.preview_json()
            elif ext in ('.xlsx', '.xls'):
                self.preview_xlsx()
            elif ext in ('.html', '.htm'):
                self.preview_html()
            elif ext == '.epub':
                self.preview_epub()
            elif ext in ('.mp3', '.wav', '.flac', '.ogg', '.m4a', '.aac', '.wma'):
                self.preview_audio()
            elif ext in ('.mp4', '.avi', '.mkv', '.mov', '.webm', '.wmv', '.flv'):
                self.preview_video()
            else:
                self.preview_unsupported()
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur lors du chargement de l\'aperçu:')}\n{e}")

    # PDF
    def preview_pdf(self):
        try:
            pdf_document = fitz.open(self.file_path)
            if len(pdf_document) == 0:
                self._show_error(self.translate_text("PDF vide"))
                return

            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            container = QWidget()
            vbox = QVBoxLayout(container)
            vbox.setSpacing(8)

            max_pages = min(len(pdf_document), 5)
            for i in range(max_pages):
                page = pdf_document.load_page(i)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                import tempfile, os as _os
                tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                tmp.close()
                pix.save(tmp.name)
                pixmap = QPixmap(tmp.name)
                _os.unlink(tmp.name)
                if not pixmap.isNull():
                    if pixmap.width() > 660:
                        pixmap = pixmap.scaledToWidth(660, Qt.SmoothTransformation)
                    lbl = QLabel()
                    lbl.setPixmap(pixmap)
                    lbl.setAlignment(Qt.AlignHCenter)
                    vbox.addWidget(lbl)

            if len(pdf_document) > max_pages:
                more = QLabel(f"… {len(pdf_document) - max_pages} {self.translate_text('pages supplémentaires')}")
                more.setAlignment(Qt.AlignCenter)
                more.setStyleSheet("color: gray; font-style: italic; padding: 6px;")
                vbox.addWidget(more)

            pdf_document.close()
            vbox.addStretch()
            scroll.setWidget(container)
            self._clear_content()
            self.content_layout.addWidget(scroll)

        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur PDF:')} {e}")

    # Image
    def preview_image(self):
        try:
            pixmap = QPixmap(self.file_path)
            if pixmap.isNull():
                self._show_error(self.translate_text("Format d'image non supporté"))
                return
            pixmap = pixmap.scaled(660, 580, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            lbl = QLabel()
            lbl.setPixmap(pixmap)
            lbl.setAlignment(Qt.AlignCenter)
            scroll.setWidget(lbl)
            self._clear_content()
            self.content_layout.addWidget(scroll)
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur image:')} {e}")

    # Word
    def preview_word(self):
        try:
            from docx import Document as _DocxDocument
            doc = _DocxDocument(self.file_path)
            from PySide6.QtWidgets import QTextEdit
            te = QTextEdit()
            te.setReadOnly(True)
            html_parts = []
            for para in doc.paragraphs[:80]:
                text = para.text
                if not text.strip():
                    html_parts.append("<br>")
                    continue
                style = para.style.name if para.style else ""
                if "Heading 1" in style:
                    html_parts.append(f"<h2>{text}</h2>")
                elif "Heading 2" in style:
                    html_parts.append(f"<h3>{text}</h3>")
                else:
                    html_parts.append(f"<p style='margin:2px 0'>{text}</p>")
            if len(doc.paragraphs) > 80:
                html_parts.append("<p style='color:gray;font-style:italic'>[...]</p>")
            te.setHtml("".join(html_parts) if html_parts else self.translate_text("Document vide"))
            self._clear_content()
            self.content_layout.addWidget(te)
        except Exception as e:
            self._show_error(f"{self.translate_text('Impossible de prévisualiser le document Word')}\n{e}")

    # Plain text / code
    def preview_text(self, syntax=""):
        try:
            from PySide6.QtWidgets import QTextEdit
            from PySide6.QtGui import QFont
            with open(self.file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read(200_000)
            te = QTextEdit()
            te.setReadOnly(True)
            te.setFont(QFont("Consolas, Courier New, monospace", 10))
            te.setPlainText(content)
            self._clear_content()
            self.content_layout.addWidget(te)
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur texte:')} {e}")

    # RTF
    def preview_rtf(self):
        name = Path(self.file_path).name
        self.preview_label.setParent(self.content_area)
        self.content_layout.addWidget(self.preview_label)
        self.preview_label.setText(
            f"🚫  {self.translate_text('Aperçu non disponible pour ce type de fichier')}\n({name})"
        )

    # CSV
    def preview_csv(self):
        try:
            import csv
            from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView
            with open(self.file_path, 'r', encoding='utf-8', errors='replace', newline='') as f:
                reader = csv.reader(f)
                rows = [r for _, r in zip(range(200), reader)]

            if not rows:
                self._show_error(self.translate_text("Fichier CSV vide"))
                return

            ncols = max(len(r) for r in rows)
            table = QTableWidget(len(rows), ncols)
            table.setHorizontalHeader
            # First row as header
            header_row = rows[0]
            table.setHorizontalHeaderLabels(header_row + [''] * (ncols - len(header_row)))
            for ri, row in enumerate(rows[1:], 0):
                for ci, val in enumerate(row):
                    table.setItem(ri, ci, QTableWidgetItem(val))
            table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
            table.horizontalHeader().setStretchLastSection(True)
            table.setEditTriggers(QTableWidget.NoEditTriggers)
            self._clear_content()
            self.content_layout.addWidget(table)
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur CSV:')} {e}")

    # JSON
    def preview_json(self):
        try:
            import json
            from PySide6.QtWidgets import QTextEdit
            from PySide6.QtGui import QFont, QColor, QTextCharFormat, QSyntaxHighlighter
            with open(self.file_path, 'r', encoding='utf-8', errors='replace') as f:
                raw = f.read(500_000)
            try:
                parsed = json.loads(raw)
                pretty = json.dumps(parsed, indent=2, ensure_ascii=False)
            except Exception:
                pretty = raw

            te = QTextEdit()
            te.setReadOnly(True)
            te.setFont(QFont("Consolas", 10))
            te.setPlainText(pretty)

            # Simple JSON syntax highlighting
            class _JsonHL(QSyntaxHighlighter):
                def highlightBlock(self, text):
                    import re
                    fmt_key   = QTextCharFormat(); fmt_key.setForeground(QColor("#569cd6"))
                    fmt_str   = QTextCharFormat(); fmt_str.setForeground(QColor("#ce9178"))
                    fmt_num   = QTextCharFormat(); fmt_num.setForeground(QColor("#b5cea8"))
                    fmt_kw    = QTextCharFormat(); fmt_kw.setForeground(QColor("#569cd6"))
                    for m in re.finditer(r'"[^"\\]*(?:\\.[^"\\]*)*"\s*:', text):
                        self.setFormat(m.start(), m.end()-m.start(), fmt_key)
                    for m in re.finditer(r'(?<!\w)"[^"\\]*(?:\\.[^"\\]*)*"', text):
                        self.setFormat(m.start(), m.end()-m.start(), fmt_str)
                    for m in re.finditer(r'\b-?\d+\.?\d*([eE][+-]?\d+)?\b', text):
                        self.setFormat(m.start(), m.end()-m.start(), fmt_num)
                    for m in re.finditer(r'\b(true|false|null)\b', text):
                        self.setFormat(m.start(), m.end()-m.start(), fmt_kw)

            _JsonHL(te.document())
            self._clear_content()
            self.content_layout.addWidget(te)
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur JSON:')} {e}")

    # XLSX
    def preview_xlsx(self):
        try:
            import openpyxl
            from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView, QTabWidget
            wb = openpyxl.load_workbook(self.file_path, read_only=True, data_only=True)
            tabs = QTabWidget()
            for sheet_name in wb.sheetnames[:6]:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows(max_row=200, values_only=True))
                if not rows:
                    continue
                ncols = max((len(r) for r in rows), default=0)
                table = QTableWidget(len(rows), ncols)
                table.setHorizontalHeaderLabels([str(i+1) for i in range(ncols)])
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        table.setItem(ri, ci, QTableWidgetItem(str(val) if val is not None else ""))
                table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
                table.horizontalHeader().setStretchLastSection(True)
                table.setEditTriggers(QTableWidget.NoEditTriggers)
                tabs.addTab(table, sheet_name)
            wb.close()
            self._clear_content()
            self.content_layout.addWidget(tabs)
        except ImportError:
            self._show_error("openpyxl non installé — pip install openpyxl")
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur XLSX:')} {e}")

    # HTML
    def preview_html(self):
        # Attempt 1: QWebEngineView (full Chromium rendering)
        try:
            from PySide6.QtWebEngineWidgets import QWebEngineView
            from PySide6.QtCore import QUrl

            view = QWebEngineView()
            view.setUrl(QUrl.fromLocalFile(str(Path(self.file_path).resolve())))
            self._clear_content()
            self.content_layout.addWidget(view)
            self.resize(900, 700)
            return
        except ImportError:
            pass  # QtWebEngine not available → try next
        except Exception as e:
            pass  # Unexpected crash → try next

        # Attempt 2: fallback to syntax-highlighted source + warning banner
        self._html_fallback_view()

    def _html_fallback_view(self):
        """
        Fallback when QtWebEngineWidgets is missing from the build:
        displays the HTML source code with full syntax highlighting
        (HTML + inline CSS in <style> + inline JS in <script>)
        and an explicit warning banner for the user.
        """
        from PySide6.QtWidgets import QTextEdit, QLabel
        from PySide6.QtGui import QFont, QColor, QTextCharFormat, QSyntaxHighlighter
        import re

        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='replace') as f:
                source = f.read(500_000)
        except Exception as e:
            self._show_error(f"Erreur lecture HTML : {e}")
            return

        # Source editor
        te = QTextEdit()
        te.setReadOnly(True)
        te.setFont(QFont("Consolas", 10))

        class _HtmlCssJsHL(QSyntaxHighlighter):
            """
            Three-level syntax highlighting:
              • HTML  — tags, attributes, values, comments, DOCTYPE
              • CSS   — selectors, properties, values, @-rules, comments
                        inside each <style>…</style> block
              • JS    — keywords, strings, numbers, comments, regex
                        inside each <script>…</script> block
            """

            # formats (instantiated once)
            def __init__(self, document, source_text=""):
                super().__init__(document)

                def _fmt(hex_color, bold=False, italic=False):
                    f = QTextCharFormat()
                    f.setForeground(QColor(hex_color))
                    if bold:   f.setFontWeight(700)
                    if italic: f.setFontItalic(True)
                    return f

                # HTML
                self._html_tag     = _fmt("#569cd6", bold=True)   # <div>, </p> …
                self._html_attr    = _fmt("#9cdcfe")               # class=, href=
                self._html_val     = _fmt("#ce9178")               # "value"
                self._html_comment = _fmt("#6a9955", italic=True)  # <!-- -->
                self._html_doctype = _fmt("#c586c0")               # <!DOCTYPE>
                self._html_entity  = _fmt("#f0a070")               # &amp; &#x27;

                # CSS
                self._css_selector = _fmt("#d7ba7d", bold=True)    # .foo, #bar, div
                self._css_prop     = _fmt("#9cdcfe")               # color:, margin:
                self._css_val      = _fmt("#ce9178")               # red, 12px, #fff
                self._css_atrule   = _fmt("#c586c0", bold=True)    # @media, @import
                self._css_comment  = _fmt("#6a9955", italic=True)  # /* … */
                self._css_number   = _fmt("#b5cea8")               # 42px, 1.5em

                # JS
                self._js_keyword   = _fmt("#569cd6", bold=True)
                self._js_builtin   = _fmt("#4ec9b0")
                self._js_string    = _fmt("#ce9178")
                self._js_template  = _fmt("#ce9178")
                self._js_number    = _fmt("#b5cea8")
                self._js_comment   = _fmt("#6a9955", italic=True)
                self._js_regex     = _fmt("#d16969")
                self._js_operator  = _fmt("#d4d4d4")

                self._JS_KEYWORDS = {
                    "var","let","const","function","class","return","if","else",
                    "for","while","do","switch","case","break","continue","new",
                    "delete","typeof","instanceof","in","of","try","catch",
                    "finally","throw","async","await","import","export","from",
                    "default","extends","super","this","static","yield","void",
                    "null","undefined","true","false","debugger","with",
                }
                self._JS_BUILTINS = {
                    "console","document","window","Array","Object","String",
                    "Number","Boolean","Math","Date","JSON","Promise","fetch",
                    "setTimeout","setInterval","clearTimeout","clearInterval",
                    "parseInt","parseFloat","isNaN","isFinite","encodeURI",
                    "decodeURI","Map","Set","WeakMap","WeakSet","Symbol",
                    "Proxy","Reflect","Error","TypeError","RangeError",
                    "localStorage","sessionStorage","navigator","location",
                    "history","screen","alert","confirm","prompt",
                }

                # Compute <style> and <script> line ranges directly from
                # source_text — BEFORE the document is populated, to avoid
                # triggering highlightBlock → toPlainText() → infinite loop.
                self._style_ranges  = []
                self._script_ranges = []
                if source_text:
                    self._compute_ranges_from_source(source_text)

            def _compute_ranges_from_source(self, text):
                """
                Computes line ranges for <style> and <script> blocks directly
                from raw text (without touching the QTextDocument).
                Called once in __init__, never from highlightBlock.
                """
                lines = text.splitlines()

                def _line_of_pos(pos):
                    return text[:pos].count("\n")

                for open_pat, close_pat, target in [
                    (r'<style[^>]*>',  r'</style\s*>',  self._style_ranges),
                    (r'<script[^>]*>', r'</script\s*>', self._script_ranges),
                ]:
                    for mo in re.finditer(open_pat, text, re.IGNORECASE):
                        inner_start = mo.end()
                        mc = re.search(close_pat, text[inner_start:], re.IGNORECASE)
                        inner_end = inner_start + mc.start() if mc else len(text)
                        target.append((_line_of_pos(inner_start), _line_of_pos(inner_end)))

            # helpers
            def _in_range(self, block_num, ranges):
                return any(s <= block_num <= e for s, e in ranges)

            def _apply(self, text, pattern, fmt, flags=0):
                for m in re.finditer(pattern, text, flags):
                    self.setFormat(m.start(), len(m.group()), fmt)

            # main entry
            def highlightBlock(self, text):
                bn = self.currentBlock().blockNumber()
                in_style  = self._in_range(bn, self._style_ranges)
                in_script = self._in_range(bn, self._script_ranges)

                if in_style:
                    self._highlight_css(text)
                elif in_script:
                    self._highlight_js(text)
                else:
                    self._highlight_html(text)

            # HTML
            def _highlight_html(self, text):
                self._apply(text, r'<!--.*?-->', self._html_comment)
                self._apply(text, r'<!DOCTYPE[^>]*>', self._html_doctype, re.IGNORECASE)
                self._apply(text, r'</?[\w\-:.]+', self._html_tag)
                self._apply(text, r'\b[\w\-:]+=', self._html_attr)   # remove trailing =
                # fix: re-colour just the name (strip the =)
                for m in re.finditer(r'\b([\w\-:]+)=', text):
                    self.setFormat(m.start(1), len(m.group(1)), self._html_attr)
                self._apply(text, r'(["\'])(?:(?!\1).)*\1', self._html_val)
                self._apply(text, r'&(?:#\d+|#x[\da-fA-F]+|[a-zA-Z]\w*);', self._html_entity)

            # CSS
            def _highlight_css(self, text):
                # /* comments */
                self._apply(text, r'/\*.*?\*/', self._css_comment)
                # @-rules
                self._apply(text, r'@[\w-]+', self._css_atrule)
                # property: value  (inside { })
                self._apply(text, r'[\w-]+(?=\s*:)', self._css_prop)
                # numbers with units
                self._apply(text, r'-?\d+\.?\d*(?:%|px|em|rem|vh|vw|vmin|vmax|pt|cm|mm|s|ms)?',
                            self._css_number)
                # strings
                self._apply(text, r'(["\'])(?:(?!\1).)*\1', self._css_val)
                # hex colors
                self._apply(text, r'#[0-9a-fA-F]{3,8}\b', self._css_val)
                # selectors (lines that don't contain : or end with {)
                if re.search(r'\{', text) or (not re.search(r':', text)):
                    self._apply(text, r'[.#]?[\w][\w-]*(?=[^:]*\{)', self._css_selector)

            # JS
            def _highlight_js(self, text):
                # single-line comments //
                self._apply(text, r'//.*', self._js_comment)
                # multi-line comments /* */
                self._apply(text, r'/\*.*?\*/', self._js_comment)
                # template literals `…`
                self._apply(text, r'`(?:[^`\\]|\\.)*`', self._js_template)
                # strings
                self._apply(text, r'(["\'])(?:(?!\1|\\).|\\.)*\1', self._js_string)
                # numbers
                self._apply(text, r'\b0[xX][\da-fA-F]+\b|\b\d+\.?\d*(?:[eE][+-]?\d+)?\b',
                            self._js_number)
                # keywords
                for m in re.finditer(r'\b([a-zA-Z_$][\w$]*)\b', text):
                    word = m.group(1)
                    if word in self._JS_KEYWORDS:
                        self.setFormat(m.start(), len(word), self._js_keyword)
                    elif word in self._JS_BUILTINS:
                        self.setFormat(m.start(), len(word), self._js_builtin)
                # simple regex literals  /pattern/flags  (heuristic)
                self._apply(text, r'(?<![<\w\d])/(?:[^/\\\n]|\\.)+/[gimsuy]*', self._js_regex)

        _HtmlCssJsHL(te.document(), source)
        te.setPlainText(source)

        # Warning banner
        warn = QLabel(self.translate_text("preview_unavailable"))
        warn.setStyleSheet(
            "background:#2d2400; color:#ffcc00; "
            "padding:6px 12px; font-size:12px; border-radius:4px;"
        )
        warn.setWordWrap(True)

        self._clear_content()
        self.content_layout.addWidget(warn)
        self.content_layout.addWidget(te)

    # EPUB
    def preview_epub(self):
        try:
            import ebooklib
            from ebooklib import epub
            from PySide6.QtWidgets import QTextEdit
            book = epub.read_epub(self.file_path)
            html_parts = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    html_parts.append(item.get_content().decode('utf-8', errors='replace'))
                    if len(html_parts) >= 3:
                        break
            if not html_parts:
                self._show_error(self.translate_text("EPUB vide ou contenu non extractible"))
                return
            combined = "<hr>".join(html_parts[:3])
            te = QTextEdit()
            te.setReadOnly(True)
            te.setHtml(combined)
            self._clear_content()
            self.content_layout.addWidget(te)
        except ImportError:
            self._show_error("ebooklib non installé — pip install ebooklib")
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur EPUB:')} {e}")

    # Audio
    def preview_audio(self):
        try:
            from PySide6.QtMultimedia import QMediaPlayer, QAudioOutput
            from PySide6.QtCore import QUrl
            from PySide6.QtWidgets import QSlider, QHBoxLayout, QVBoxLayout, QWidget, QLabel, QPushButton
            from PySide6.QtCore import Qt

            player = QMediaPlayer()
            audio_out = QAudioOutput()
            player.setAudioOutput(audio_out)
            audio_out.setVolume(0.8)
            player.setSource(QUrl.fromLocalFile(str(Path(self.file_path).resolve())))
            self._media_player = player
            self._audio_output = audio_out  # keep alive

            container = QWidget()
            vbox = QVBoxLayout(container)
            vbox.setSpacing(12)

            # File info
            info = QLabel(f"🎵  <b>{Path(self.file_path).name}</b>")
            info.setStyleSheet("font-size: 14px; padding: 8px;")
            info.setAlignment(Qt.AlignCenter)
            vbox.addWidget(info)

            # Time display
            time_lbl = QLabel("0:00 / 0:00")
            time_lbl.setAlignment(Qt.AlignCenter)
            time_lbl.setStyleSheet("font-size: 12px; color: gray;")
            vbox.addWidget(time_lbl)

            # Seek slider
            seek_slider = QSlider(Qt.Horizontal)
            seek_slider.setRange(0, 0)
            vbox.addWidget(seek_slider)

            # Controls row
            ctrl = QWidget()
            hbox = QHBoxLayout(ctrl)
            hbox.setSpacing(12)

            play_btn  = QPushButton("▶  " + self.translate_text("Lire"))
            pause_btn = QPushButton("⏸  " + self.translate_text("Pause"))
            stop_btn  = QPushButton("⏹  " + self.translate_text("Arrêter"))
            for btn in (play_btn, pause_btn, stop_btn):
                btn.setFixedHeight(34)
                hbox.addWidget(btn)

            # Volume
            vol_lbl = QLabel("🔊")
            vol_slider = QSlider(Qt.Horizontal)
            vol_slider.setRange(0, 100)
            vol_slider.setValue(80)
            vol_slider.setFixedWidth(90)
            hbox.addStretch()
            hbox.addWidget(vol_lbl)
            hbox.addWidget(vol_slider)
            vbox.addWidget(ctrl)

            # Connections
            play_btn.clicked.connect(player.play)
            pause_btn.clicked.connect(player.pause)
            stop_btn.clicked.connect(player.stop)
            vol_slider.valueChanged.connect(lambda v: audio_out.setVolume(v / 100))

            def _on_duration(dur):
                seek_slider.setRange(0, dur)
                secs = dur // 1000
                time_lbl.setText(f"0:00 / {secs//60}:{secs%60:02d}")

            def _on_position(pos):
                seek_slider.blockSignals(True)
                seek_slider.setValue(pos)
                seek_slider.blockSignals(False)
                dur = player.duration()
                def _fmt(ms):
                    s = ms // 1000
                    return f"{s//60}:{s%60:02d}"
                time_lbl.setText(f"{_fmt(pos)} / {_fmt(dur)}")

            def _on_seek(val):
                player.setPosition(val)

            player.durationChanged.connect(_on_duration)
            player.positionChanged.connect(_on_position)
            seek_slider.sliderMoved.connect(_on_seek)

            vbox.addStretch()
            self._clear_content()
            self.content_layout.addWidget(container)

        except ImportError as e:
            self._show_error(f"PySide6.QtMultimedia manquant: {e}")
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur audio:')} {e}")

    # Video
    def preview_video(self):
        try:
            from PySide6.QtMultimedia import QMediaPlayer, QAudioOutput
            from PySide6.QtMultimediaWidgets import QVideoWidget
            from PySide6.QtCore import QUrl
            from PySide6.QtWidgets import QSlider, QHBoxLayout, QVBoxLayout, QWidget, QLabel, QPushButton
            from PySide6.QtCore import Qt

            player = QMediaPlayer()
            audio_out = QAudioOutput()
            player.setAudioOutput(audio_out)
            audio_out.setVolume(0.8)
            self._media_player = player
            self._audio_output = audio_out

            video_widget = QVideoWidget()
            video_widget.setMinimumHeight(300)
            player.setVideoOutput(video_widget)
            player.setSource(QUrl.fromLocalFile(str(Path(self.file_path).resolve())))

            container = QWidget()
            vbox = QVBoxLayout(container)
            vbox.setSpacing(6)
            vbox.addWidget(video_widget, 1)

            # Time + seek
            time_lbl = QLabel("0:00 / 0:00")
            time_lbl.setAlignment(Qt.AlignCenter)
            time_lbl.setStyleSheet("font-size: 11px; color: gray;")
            seek_slider = QSlider(Qt.Horizontal)
            seek_slider.setRange(0, 0)
            vbox.addWidget(seek_slider)
            vbox.addWidget(time_lbl)

            # Controls
            ctrl = QWidget()
            hbox = QHBoxLayout(ctrl)
            hbox.setSpacing(10)

            play_btn  = QPushButton("▶  " + self.translate_text("Lire"))
            pause_btn = QPushButton("⏸  " + self.translate_text("Pause"))
            stop_btn  = QPushButton("⏹  " + self.translate_text("Arrêter"))
            for btn in (play_btn, pause_btn, stop_btn):
                btn.setFixedHeight(32)
                hbox.addWidget(btn)

            vol_lbl = QLabel("🔊")
            vol_slider = QSlider(Qt.Horizontal)
            vol_slider.setRange(0, 100)
            vol_slider.setValue(80)
            vol_slider.setFixedWidth(80)
            hbox.addStretch()
            hbox.addWidget(vol_lbl)
            hbox.addWidget(vol_slider)
            vbox.addWidget(ctrl)

            # Connections
            play_btn.clicked.connect(player.play)
            pause_btn.clicked.connect(player.pause)
            stop_btn.clicked.connect(player.stop)
            vol_slider.valueChanged.connect(lambda v: audio_out.setVolume(v / 100))

            def _on_duration(dur):
                seek_slider.setRange(0, dur)

            def _on_position(pos):
                seek_slider.blockSignals(True)
                seek_slider.setValue(pos)
                seek_slider.blockSignals(False)
                def _fmt(ms):
                    s = ms // 1000
                    return f"{s//60}:{s%60:02d}"
                time_lbl.setText(f"{_fmt(pos)} / {_fmt(player.duration())}")

            player.durationChanged.connect(_on_duration)
            player.positionChanged.connect(_on_position)
            seek_slider.sliderMoved.connect(player.setPosition)

            self._clear_content()
            self.content_layout.addWidget(container)
            self.resize(900, 620)

        except ImportError as e:
            self._show_error(f"PySide6.QtMultimediaWidgets manquant: {e}")
        except Exception as e:
            self._show_error(f"{self.translate_text('Erreur vidéo:')} {e}")

    # Unsupported
    def preview_unsupported(self):
        ext = Path(self.file_path).suffix.upper() or "?"
        self.preview_label.setParent(self.content_area)
        self.content_layout.addWidget(self.preview_label)
        self.preview_label.setText(
            f"🚫  {self.translate_text('Aperçu non disponible pour ce type de fichier')}\n({ext})"
        )

    def closeEvent(self, event):
        if self._media_player:
            try:
                self._media_player.stop()
            except Exception:
                pass
        super().closeEvent(event)

class PasswordDialog(QDialog):
    def __init__(self, parent=None, language="fr"):
        super().__init__(parent)
        self.language = language
        self._tm = _make_tm(language)
        self.setWindowTitle(self.translate_text("Protéger PDF avec mot de passe"))
        self.setup_ui()

    def setup_ui(self):
        layout = QFormLayout(self)

        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.Password)
        self.confirm_input = QLineEdit()
        self.confirm_input.setEchoMode(QLineEdit.Password)

        layout.addRow(self.translate_text("Mot de passe:"), self.password_input)
        layout.addRow(self.translate_text("Confirmer mot de passe:"), self.confirm_input)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        ok_button = buttons.button(QDialogButtonBox.Ok)
        if ok_button:
            ok_button.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
                QPushButton:pressed {
                    background-color: #1e7e34;
                }
            """)
        cancel_button = buttons.button(QDialogButtonBox.Cancel)
        if cancel_button:
            cancel_button.setStyleSheet("""
                QPushButton {
                    background-color: #B55454;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #A04040;
                }
                QPushButton:pressed {
                    background-color: #8B3030;
                }
            """)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout.addRow(buttons)

    def translate_text(self, text):
        return self._tm.translate_text(text)

    def get_password(self):
        return self.password_input.text()

class SplitDialog(QDialog):

    def __init__(self, total_pages, parent=None, language="fr"):
        super().__init__(parent)
        self.total_pages = total_pages
        self.language = language
        self._tm = _make_tm(language)
        self.setWindowTitle(self.translate_text("Diviser PDF"))
        self.setup_ui()

    def setup_ui(self):
        layout = QFormLayout(self)

        self.split_method = QComboBox()
        self.split_method.addItems([
            self.translate_text("Par pages"),
            self.translate_text("Toutes les pages"),
            self.translate_text("Plage de pages")
        ])

        self.page_interval = QSpinBox()
        self.page_interval.setMinimum(1)
        self.page_interval.setMaximum(self.total_pages)
        self.page_interval.setValue(1)

        self.start_page = QSpinBox()
        self.start_page.setMinimum(1)
        self.start_page.setMaximum(self.total_pages)
        self.start_page.setValue(1)

        self.end_page = QSpinBox()
        self.end_page.setMinimum(1)
        self.end_page.setMaximum(self.total_pages)
        self.end_page.setValue(self.total_pages)

        layout.addRow(self.translate_text("Méthode de division:"), self.split_method)
        layout.addRow(self.translate_text("Intervalle (pages):"), self.page_interval)
        layout.addRow(self.translate_text("Page de début:"), self.start_page)
        layout.addRow(self.translate_text("Page de fin:"), self.end_page)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        ok_button = buttons.button(QDialogButtonBox.Ok)
        if ok_button:
            ok_button.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
                QPushButton:pressed {
                    background-color: #1e7e34;
                }
            """)
        cancel_button = buttons.button(QDialogButtonBox.Cancel)
        if cancel_button:
            cancel_button.setStyleSheet("""
                QPushButton {
                    background-color: #B55454;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #A04040;
                }
                QPushButton:pressed {
                    background-color: #8B3030;
                }
            """)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout.addRow(buttons)

    def translate_text(self, text):
        return self._tm.translate_text(text)

class CompressionDialog(QDialog):
    def __init__(self, parent=None, language="fr"):
        super().__init__(parent)
        self.language = language
        self._tm = _make_tm(language)
        self.setWindowTitle(self.translate_text("Compresser des Fichiers"))
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)

        # Selection information
        selection_info = QLabel()
        selection_info.setStyleSheet("font-weight: bold; color: #4dabf7;")
        selection_info.setWordWrap(True)
        layout.addWidget(selection_info)

        # Archive format
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel(self.translate_text("Format d'archive:")))
        self.format_combo = QComboBox()
        self.format_combo.addItems(["ZIP", "RAR", "TAR.GZ", "TAR"])
        self.format_combo.currentIndexChanged.connect(self.update_format_info)
        format_layout.addWidget(self.format_combo)
        layout.addLayout(format_layout)

        # Compression level
        level_layout = QHBoxLayout()
        level_layout.addWidget(QLabel(self.translate_text("Niveau de compression:")))
        self.compression_level = QComboBox()
        self.compression_level.addItems([
            self.translate_text("Normal"),
            self.translate_text("Haute compression"),
            self.translate_text("Compression maximale")
        ])
        level_layout.addWidget(self.compression_level)
        layout.addLayout(level_layout)

        # Archive name
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel(self.translate_text("Nom de l'archive:")))
        self.filename_input = QLineEdit()
        self.filename_input.setText("archive_compressee")
        name_layout.addWidget(self.filename_input)
        layout.addLayout(name_layout)

        # Options
        options_group = QGroupBox(self.translate_text("Options"))
        options_layout = QVBoxLayout(options_group)

        self.encryption_checkbox = AnimatedCheckBox(self.translate_text("Protéger par mot de passe"))
        self.encryption_checkbox.stateChanged.connect(self.on_encryption_changed)

        self.split_checkbox = AnimatedCheckBox(self.translate_text("Fractionner l'archive en plusieurs parties"))
        self.split_checkbox.setChecked(False)
        self.split_checkbox.stateChanged.connect(self.on_split_changed)

        self.delete_original_checkbox = AnimatedCheckBox(self.translate_text("Supprimer les fichiers originaux après compression"))

        split_size_label = QLabel(self.translate_text("Taille par partie:"))

        self.split_size_spin = QSpinBox()
        self.split_size_spin.setRange(1, 10000)
        self.split_size_spin.setValue(100)
        self.split_size_spin.setSuffix(" MB")
        self.split_size_spin.setEnabled(False)
        self.split_size_spin.valueChanged.connect(self.on_split_size_changed)

        self.split_preset_combo = QComboBox()
        self.split_preset_combo.addItems([
            "10 MB (Email)",
            "25 MB (Email)",
            self.translate_text("100 MB (Partage web)"),
            "700 MB (CD)",
            "4.7 GB (DVD)",
            self.translate_text("8.5 GB (Double couche DVD)"),
            self.translate_text("Personnalisé")
        ])
        self.split_preset_combo.setEnabled(False)
        self.split_preset_combo.setCurrentIndex(2)

        split_size_layout = QHBoxLayout()
        split_size_layout.addWidget(QLabel(self.translate_text("Taille par partie:")))
        split_size_layout.addWidget(self.split_size_spin)
        split_size_layout.addWidget(self.split_preset_combo)
        split_size_layout.addStretch()

        self.split_info_label = QLabel()
        self.split_info_label.setStyleSheet("font-size: 10px; color: #6c757d; font-style: italic;")
        self.split_info_label.setWordWrap(True)
        self.split_info_label.setVisible(False)

        options_layout.addWidget(self.encryption_checkbox)
        options_layout.addWidget(self.split_checkbox)
        options_layout.addWidget(self.delete_original_checkbox)
        options_layout.addLayout(split_size_layout)
        options_layout.addWidget(self.split_info_label)

        layout.addWidget(options_group)

        self.format_info = QLabel()
        self.format_info.setStyleSheet("font-size: 11px; color: #6c757d; font-style: italic;")
        self.format_info.setWordWrap(True)
        layout.addWidget(self.format_info)

        self.split_note_label = QLabel()
        self.split_note_label.setStyleSheet("font-size: 10px; color: #dc3545; font-style: italic;")
        self.split_note_label.setWordWrap(True)
        self.split_note_label.setText(self.translate_text(
            "Note: Le fractionnement n'est disponible que pour les formats ZIP et RAR. "
            "Les archives TAR/TAR.GZ ne peuvent pas être fractionnées."
        ))
        self.split_note_label.setVisible(False)
        layout.addWidget(self.split_note_label)

        # Buttons
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        ok_button = buttons.button(QDialogButtonBox.Ok)
        if ok_button:
            ok_button.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
                QPushButton:pressed {
                    background-color: #1e7e34;
                }
            """)
        cancel_button = buttons.button(QDialogButtonBox.Cancel)
        if cancel_button:
            cancel_button.setStyleSheet("""
                QPushButton {
                    background-color: #B55454;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #A04040;
                }
                QPushButton:pressed {
                    background-color: #8B3030;
                }
            """)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        # Connect signals
        self.split_checkbox.stateChanged.connect(self.on_split_changed)
        self.format_combo.currentIndexChanged.connect(self.on_format_changed)
        self.split_preset_combo.currentIndexChanged.connect(self.on_split_preset_changed)
        self.split_size_spin.valueChanged.connect(lambda: self.on_split_preset_changed(0))
        self.update_split_controls_state()

        self.update_format_info()
        self.on_format_changed()

        self.apply_theme_style()

    def update_split_controls_state(self):
        """Update the state of all splitting controls"""
        is_split_enabled = self.split_checkbox.isChecked()
        is_custom = self.split_preset_combo.currentText() == self.translate_text("Personnalisé")

        self.split_preset_combo.setEnabled(is_split_enabled)
        self.split_size_spin.setEnabled(is_split_enabled and is_custom)

        if is_split_enabled:
            self.update_split_info_message()

    def apply_theme_style(self):
        """Apply style according to the parent application theme"""
        parent = self.parent()
        if hasattr(parent, 'dark_mode'):
            if parent.dark_mode:
                self.setStyleSheet("""
                    QDialog { background-color: #2d333b; }
                    QLabel { color: #e9ecef; }
                    QGroupBox {
                        color: #adb5bd;
                        border: 1px solid #495057;
                        border-radius: 6px;
                        margin-top: 10px;
                    }
                    QGroupBox::title {
                        subcontrol-origin: margin;
                        left: 10px;
                        padding: 0 5px;
                    }
                    QComboBox, QLineEdit, QSpinBox {
                        background-color: #495057;
                        color: #e9ecef;
                        border: 1px solid #6c757d;
                        border-radius: 4px;
                        padding: 5px;
                    }
                    QCheckBox { color: #e9ecef; spacing: 8px; }
                    QCheckBox::indicator {
                        width: 16px; height: 16px;
                        border-radius: 3px;
                        border: 2px solid #6c757d;
                    }
                    QCheckBox::indicator:checked {
                        background-color: #4dabf7;
                        border: 2px solid #4dabf7;
                    }
                """)
            else:
                self.setStyleSheet("""
                    QDialog { background-color: #f8f9fa; }
                    QLabel { color: #212529; }
                    QGroupBox {
                        color: #495057;
                        border: 1px solid #dee2e6;
                        border-radius: 6px;
                        margin-top: 10px;
                    }
                    QGroupBox::title {
                        subcontrol-origin: margin;
                        left: 10px;
                        padding: 0 5px;
                    }
                    QComboBox, QLineEdit, QSpinBox {
                        background-color: #ffffff;
                        color: #212529;
                        border: 1px solid #ced4da;
                        border-radius: 4px;
                        padding: 5px;
                    }
                    QCheckBox { color: #212529; spacing: 8px; }
                    QCheckBox::indicator {
                        width: 16px; height: 16px;
                        border-radius: 3px;
                        border: 2px solid #adb5bd;
                    }
                    QCheckBox::indicator:checked {
                        background-color: #4dabf7;
                        border: 2px solid #4dabf7;
                    }
                """)

    def on_split_size_changed(self, value):
        if self.split_checkbox.isChecked() and self.split_preset_combo.currentText() == self.translate_text("Personnalisé"):
            template = self.translate_text("split_info_custom")
            message = template.format(value)
            self.split_info_label.setText(message)
            self.split_info_label.setVisible(True)
            print(f"[DEBUG] Custom size changed: {value} MB")

    def on_split_changed(self, state):
        is_enabled = state == Qt.Checked
        self.update_split_controls_state()

    def on_split_preset_changed(self, index):
        preset_text = self.split_preset_combo.currentText()

        if preset_text == self.translate_text("Personnalisé"):
            self.split_size_spin.setEnabled(True)
        else:
            self.split_size_spin.setEnabled(False)

            preset_sizes = {
                "10 MB (Email)": 10,
                "25 MB (Email)": 25,
                "100 MB (Partage web)": 100,
                "700 MB (CD)": 700,
                "4.7 GB (DVD)": 4700,
                "8.5 GB (Double couche DVD)": 8500
            }

            if preset_text in preset_sizes:
                self.split_size_spin.setValue(preset_sizes[preset_text])

        self.update_split_info_message()

    def update_split_info_message(self):
        if self.split_checkbox.isChecked():
            preset_text = self.split_preset_combo.currentText()
            size_value = self.split_size_spin.value()

            if preset_text == self.translate_text("Personnalisé"):
                template = self.translate_text("split_info_custom")
                message = template.format(size_value)
            else:
                size_name = preset_text.split(' ')[0]
                template = self.translate_text("split_info_preset")
                message = template.format(size_value, size_name)

            self.split_info_label.setText(message)
            self.split_info_label.setVisible(True)
        else:
            self.split_info_label.setVisible(False)

    def on_format_changed(self):
        current_format = self.format_combo.currentText()

        if current_format in ["TAR", "TAR.GZ"]:
            self.split_checkbox.setEnabled(False)
            self.split_checkbox.setChecked(False)
            self.split_note_label.setVisible(True)
            self.split_size_spin.setEnabled(False)
            self.split_preset_combo.setEnabled(False)
        else:
            self.split_checkbox.setEnabled(True)
            self.split_note_label.setVisible(False)

            if self.split_checkbox.isChecked():
                self.split_preset_combo.setEnabled(True)
                if self.split_preset_combo.currentText() == self.translate_text("Personnalisé"):
                    self.split_size_spin.setEnabled(True)
                else:
                    self.split_size_spin.setEnabled(False)

                if self.split_preset_combo.currentIndex() > 0:
                    preset_text = self.split_preset_combo.currentText()
                    size_name = preset_text.split(' ')[0]
                    size_value = self.split_size_spin.value()
                    template = self.translate_text("split_info_preset")
                    self.split_info_label.setText(template.format(size_value, size_name))

    def on_encryption_changed(self, state):
        if state == Qt.Checked:
            parent = self.parent()
            if hasattr(parent, 'files_list'):
                total_size = sum(os.path.getsize(f) for f in parent.files_list if os.path.exists(f))
                total_size_mb = total_size / (1024 * 1024)

                if total_size_mb > 100:
                    message = self.translate_text(
                        f"⚠️ Attention: Vous allez chiffrer {total_size_mb:.1f} Mo de données.\n"
                        "Le chiffrement peut ralentir la compression."
                    )
                else:
                    message = self.translate_text(
                        f"Le chiffrement ZIP utilise le standard AES-256 avec pyzipper (si installé).\n\n"
                        "Pour une meilleure sécurité, assurez-vous d'avoir installé pyzipper:\n"
                        "pip install pyzipper"
                    )
            else:
                message = self.translate_text(
                    "Le chiffrement ZIP utilise le standard AES-256 avec pyzipper (si installé).\n\n"
                    "Pour une meilleure sécurité, assurez-vous d'avoir installé pyzipper:\n"
                    "pip install pyzipper"
                )

            QMessageBox.information(self,
                                    self.translate_text("Information sur le chiffrement"),
                                    message)

    def update_format_info(self):
        format_info = {
            "ZIP": self.translate_text("Format universel, compatible avec tous les systèmes. Supporte le chiffrement et le fractionnement."),
            "RAR": self.translate_text("Meilleure compression mais nécessite WinRAR/7-Zip pour décompresser. Supporte le fractionnement."),
            "TAR.GZ": self.translate_text("Standard Unix/Linux, bonne compression. Ne supporte PAS le fractionnement."),
            "TAR": self.translate_text("Archive non compressée, préférable pour sauvegarde rapide. Ne supporte PAS le fractionnement.")
        }

        current_format = self.format_combo.currentText()
        self.format_info.setText(format_info.get(current_format, ""))

        if current_format in ["TAR", "TAR.GZ"]:
            self.split_note_label.setVisible(True)
        else:
            self.split_note_label.setVisible(False)

    def get_compression_settings(self):
        is_split_enabled = self.split_checkbox.isChecked()
        current_format = self.format_combo.currentText()
        is_split_supported = current_format in ["ZIP", "RAR",
                                                self.translate_text("ZIP"),
                                                self.translate_text("RAR")]

        print(f"[DEBUG GET_SETTINGS] Split checkbox: {is_split_enabled}")
        print(f"[DEBUG GET_SETTINGS] Format: {current_format}")
        print(f"[DEBUG GET_SETTINGS] Split supported: {is_split_supported}")
        print(f"[DEBUG GET_SETTINGS] Split size: {self.split_size_spin.value()}")

        split_size = 0
        if is_split_enabled and is_split_supported:
            split_size = self.split_size_spin.value()
            print(f"[DEBUG] Splitting size chosen: {split_size}MB")
            if split_size < 1:
                split_size = 100

        settings = {
            'format': current_format,
            'level': self.compression_level.currentText(),
            'name': self.filename_input.text().strip() or "archive_compressee",
            'password': self.encryption_checkbox.isChecked(),
            'split': is_split_enabled and is_split_supported,
            'split_size': split_size,
            'delete_originals': self.delete_original_checkbox.isChecked()
        }

        print(f"[DEBUG FINAL SETTINGS] Format: {settings['format']}")
        print(f"[DEBUG FINAL SETTINGS] Split: {settings['split']}, Size: {settings['split_size']}MB")

        return settings

    def translate_text(self, text):
        return self._tm.translate_text(text)

class BatchConvertDialog(QDialog):
    def __init__(self, parent=None, language="fr"):
        super().__init__(parent)
        self.language = language
        self._tm = _make_tm(language)
        self.setWindowTitle(self.translate_text("Conversion par Lot"))
        self.setup_ui()

    def setup_ui(self):
        layout = QFormLayout(self)

        self.format_combo = QComboBox()
        self.format_combo.addItems([
            self.translate_text("PDF"),
            self.translate_text("DOCX"),
            self.translate_text("Images PNG")
        ])

        self.quality_combo = QComboBox()
        self.quality_combo.addItems([
            self.translate_text("Haute qualité"),
            self.translate_text("Qualité standard"),
            self.translate_text("Compressé")
        ])

        layout.addRow(self.translate_text("Format cible:"), self.format_combo)
        layout.addRow(self.translate_text("Qualité:"), self.quality_combo)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        ok_button = buttons.button(QDialogButtonBox.Ok)
        if ok_button:
            ok_button.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
                QPushButton:pressed {
                    background-color: #1e7e34;
                }
            """)
        cancel_button = buttons.button(QDialogButtonBox.Cancel)
        if cancel_button:
            cancel_button.setStyleSheet("""
                QPushButton {
                    background-color: #B55454;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 6px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #A04040;
                }
                QPushButton:pressed {
                    background-color: #8B3030;
                }
            """)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addRow(buttons)

    def translate_text(self, text):
        return self._tm.translate_text(text)

class BatchRenameDialog(QDialog):
    """
    Advanced batch rename dialog.
    Features: flexible template, numbering options, case/cleanup, find/replace, live preview.
    """

    def __init__(self, files, parent=None, language="fr"):
        super().__init__(parent)
        self.files = list(files)
        self._tm = _make_tm(language)
        self.setWindowTitle(self.tr_("Renommage par Lot"))
        self.setMinimumSize(740, 600)
        self._setup_ui()
        self._refresh_preview()

    # i18n
    def tr_(self, text):
        return self._tm.translate_text(text)

    # UI
    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setSpacing(10)
        root.setContentsMargins(14, 14, 14, 14)

        # 1. Template
        tpl_box = QGroupBox(self.tr_("br_template_title"))
        tpl_lay = QVBoxLayout(tpl_box)

        hint = QLabel(self.tr_("br_template_hint"))
        hint.setStyleSheet("color: rgba(160,160,160,0.85); font-size: 11px;")
        hint.setWordWrap(True)
        tpl_lay.addWidget(hint)

        tpl_row = QHBoxLayout()
        self.tpl_input = QLineEdit("{original}")
        self.tpl_input.setPlaceholderText("{original}_{num}")
        self.tpl_input.textChanged.connect(self._refresh_preview)
        tpl_row.addWidget(self.tpl_input)

        for var in ["{original}", "{num}", "{date}", "{ext}"]:
            btn = QPushButton(var)
            btn.setFixedHeight(28)
            btn.setStyleSheet("""
                QPushButton { background: rgba(170,100,255,0.15); color: rgb(170,100,255);
                              border: 1px solid rgba(170,100,255,0.35); border-radius:5px;
                              font-size:11px; padding: 0 6px; }
                QPushButton:hover { background: rgba(170,100,255,0.28); }
            """)
            btn.clicked.connect(lambda _, v=var: self._insert_var(v))
            tpl_row.addWidget(btn)
        tpl_lay.addLayout(tpl_row)
        root.addWidget(tpl_box)

        # 2. Numbering
        num_box = QGroupBox(self.tr_("br_numbering_title"))
        num_grid = QGridLayout(num_box)
        num_grid.setColumnStretch(1, 1)
        num_grid.setColumnStretch(3, 1)

        num_grid.addWidget(QLabel(self.tr_("br_start")), 0, 0)
        self.start_spin = QSpinBox()
        self.start_spin.setRange(0, 99999)
        self.start_spin.setValue(1)
        self.start_spin.valueChanged.connect(self._refresh_preview)
        num_grid.addWidget(self.start_spin, 0, 1)

        num_grid.addWidget(QLabel(self.tr_("br_step")), 0, 2)
        self.step_spin = QSpinBox()
        self.step_spin.setRange(1, 100)
        self.step_spin.setValue(1)
        self.step_spin.valueChanged.connect(self._refresh_preview)
        num_grid.addWidget(self.step_spin, 0, 3)

        num_grid.addWidget(QLabel(self.tr_("br_padding")), 1, 0)
        self.pad_spin = QSpinBox()
        self.pad_spin.setRange(1, 6)
        self.pad_spin.setValue(3)
        self.pad_spin.valueChanged.connect(self._refresh_preview)
        num_grid.addWidget(self.pad_spin, 1, 1)

        num_grid.addWidget(QLabel(self.tr_("br_order")), 1, 2)
        self.order_combo = QComboBox()
        self.order_combo.addItems([
            self.tr_("br_order_current"),
            self.tr_("Alphabétique (A→Z)"),
            self.tr_("Alphabétique (Z→A)"),
            self.tr_("Numérique (1→9)"),
            self.tr_("Numérique (9→1)"),
            self.tr_("Date (ancien→nouveau)"),
            self.tr_("Date (nouveau→ancien)"),
        ])
        self.order_combo.currentIndexChanged.connect(self._refresh_preview)
        num_grid.addWidget(self.order_combo, 1, 3)
        root.addWidget(num_box)

        # 3. Casse & Nettoyage
        case_box = QGroupBox(self.tr_("br_case_title"))
        case_lay = QHBoxLayout(case_box)

        case_lay.addWidget(QLabel(self.tr_("br_case_label")))
        self.case_combo = QComboBox()
        self.case_combo.addItems([
            self.tr_("br_case_unchanged"),
            self.tr_("br_case_upper"),
            self.tr_("br_case_lower"),
            self.tr_("br_case_title"),
        ])
        self.case_combo.currentIndexChanged.connect(self._refresh_preview)
        case_lay.addWidget(self.case_combo)
        case_lay.addSpacing(16)

        from widgets import AnimatedCheckBox
        self.spaces_check = AnimatedCheckBox(self.tr_("br_spaces"))
        self.spaces_check.stateChanged.connect(self._refresh_preview)
        case_lay.addWidget(self.spaces_check)

        self.trim_check = AnimatedCheckBox(self.tr_("br_trim"))
        self.trim_check.stateChanged.connect(self._refresh_preview)
        case_lay.addWidget(self.trim_check)

        self.special_check = AnimatedCheckBox(self.tr_("br_special"))
        self.special_check.stateChanged.connect(self._refresh_preview)
        case_lay.addWidget(self.special_check)
        case_lay.addStretch()
        root.addWidget(case_box)

        # 4. Chercher / Remplacer
        fr_box = QGroupBox(self.tr_("br_findreplace_title"))
        fr_lay = QHBoxLayout(fr_box)
        fr_lay.addWidget(QLabel(self.tr_("br_find")))
        self.find_input = QLineEdit()
        self.find_input.setPlaceholderText(self.tr_("br_find_placeholder"))
        self.find_input.textChanged.connect(self._refresh_preview)
        fr_lay.addWidget(self.find_input)
        fr_lay.addWidget(QLabel(self.tr_("br_replace")))
        self.replace_input = QLineEdit()
        self.replace_input.setPlaceholderText(self.tr_("br_replace_placeholder"))
        self.replace_input.textChanged.connect(self._refresh_preview)
        fr_lay.addWidget(self.replace_input)
        root.addWidget(fr_box)

        # 5. Preview
        prev_box = QGroupBox(self.tr_("br_preview_title"))
        prev_lay = QVBoxLayout(prev_box)

        self.preview_table = QTableWidget(0, 2)
        self.preview_table.setHorizontalHeaderLabels([
            self.tr_("br_col_before"), self.tr_("br_col_after")
        ])
        self.preview_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.preview_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.preview_table.setAlternatingRowColors(True)
        self.preview_table.setMinimumHeight(160)
        self.preview_table.setStyleSheet(
            "QTableWidget { border: none; font-size: 12px; }"
            "QHeaderView::section { font-weight: 700; padding: 4px; font-size: 12px; }"
        )
        prev_lay.addWidget(self.preview_table)
        root.addWidget(prev_box)
        root.setStretchFactor(prev_box, 1)

        # Boutons
        btn_row = QHBoxLayout()
        self.rename_btn = QPushButton("✅ " + self.tr_("br_apply"))
        self.rename_btn.setMinimumHeight(36)
        self.rename_btn.setStyleSheet("""
            QPushButton { background:#28a745; color:white; border:none;
                          border-radius:7px; font-weight:bold; padding:6px 20px; }
            QPushButton:hover { background:#218838; }
            QPushButton:pressed { background:#1e7e34; }
        """)
        cancel_btn = QPushButton(self.tr_("Annuler"))
        cancel_btn.setMinimumHeight(36)
        cancel_btn.setStyleSheet("""
            QPushButton { background:#B55454; color:white; border:none;
                          border-radius:7px; font-weight:bold; padding:6px 20px; }
            QPushButton:hover { background:#A04040; }
            QPushButton:pressed { background:#8B3030; }
        """)
        self.rename_btn.clicked.connect(self.accept)
        cancel_btn.clicked.connect(self.reject)
        btn_row.addStretch()
        btn_row.addWidget(cancel_btn)
        btn_row.addWidget(self.rename_btn)
        root.addLayout(btn_row)

    # Helpers
    def _insert_var(self, var):
        pos = self.tpl_input.cursorPosition()
        txt = self.tpl_input.text()
        self.tpl_input.setText(txt[:pos] + var + txt[pos:])
        self.tpl_input.setCursorPosition(pos + len(var))

    def _sorted_files(self):
        import re as _re2
        files = list(self.files)
        idx = self.order_combo.currentIndex()
        if idx == 1:
            files.sort(key=lambda f: _os.path.basename(f).lower())
        elif idx == 2:
            files.sort(key=lambda f: _os.path.basename(f).lower(), reverse=True)
        elif idx == 3:
            def _nk(f):
                nums = _re2.findall(r'\d+', _os.path.splitext(_os.path.basename(f))[0])
                return [int(n) for n in nums] if nums else [0]
            files.sort(key=_nk)
        elif idx == 4:
            def _nkd(f):
                nums = _re2.findall(r'\d+', _os.path.splitext(_os.path.basename(f))[0])
                return [int(n) for n in nums] if nums else [0]
            files.sort(key=_nkd, reverse=True)
        elif idx == 5:
            files.sort(key=lambda f: _os.path.getmtime(f) if _os.path.exists(f) else 0)
        elif idx == 6:
            files.sort(key=lambda f: _os.path.getmtime(f) if _os.path.exists(f) else 0, reverse=True)
        return files

    def _apply_transforms(self, stem):
        import re as _re2
        if self.trim_check.isChecked():
            stem = stem.strip()
        if self.spaces_check.isChecked():
            stem = stem.replace(" ", "_")
        if self.special_check.isChecked():
            stem = _re2.sub(r'[^\w\-\.]', '', stem)
        find = self.find_input.text()
        if find:
            stem = stem.replace(find, self.replace_input.text())
        case_idx = self.case_combo.currentIndex()
        if case_idx == 1:
            stem = stem.upper()
        elif case_idx == 2:
            stem = stem.lower()
        elif case_idx == 3:
            stem = stem.title()
        return stem

    def _compute_new_name(self, file_path, index):
        from datetime import datetime as _dt
        stem = _os.path.splitext(_os.path.basename(file_path))[0]
        ext  = _os.path.splitext(file_path)[1]
        num  = self.start_spin.value() + index * self.step_spin.value()
        pad  = self.pad_spin.value()
        date = _dt.now().strftime("%Y-%m-%d")

        tpl = self.tpl_input.text() or "{original}"
        new_stem = (tpl
                    .replace("{original}", stem)
                    .replace("{num}",      str(num).zfill(pad))
                    .replace("{date}",     date)
                    .replace("{ext}",      ext.lstrip(".")))
        new_stem = self._apply_transforms(new_stem)
        return new_stem + ext

    # Live preview
    def _refresh_preview(self):
        from PySide6.QtGui import QColor
        files = self._sorted_files()
        self.preview_table.setRowCount(len(files))
        for i, fp in enumerate(files):
            old_name = _os.path.basename(fp)
            new_name = self._compute_new_name(fp, i)
            old_item = QTableWidgetItem(old_name)
            new_item = QTableWidgetItem(new_name)
            if old_name != new_name:
                new_item.setForeground(QColor("#4ade80"))
            self.preview_table.setItem(i, 0, old_item)
            self.preview_table.setItem(i, 1, new_item)

    # Public API
    def get_rename_plan(self):
        """Return list of (old_path, new_name) in chosen order."""
        files = self._sorted_files()
        return [(fp, self._compute_new_name(fp, i)) for i, fp in enumerate(files)]

class SettingsDialog(QDialog):
    def __init__(self, config, parent=None, language="fr"):
        super().__init__(parent)
        self.config = config
        self.language = language
        self._tm = _make_tm(language)
        self.setWindowTitle(self.translate_text("Paramètres de l'application"))
        self.setModal(True)
        self.setMinimumSize(650, 550)
        self.setMaximumHeight(700)
        self.setup_ui()

    def setup_ui(self):

        ok_button_style = """
        QPushButton {
            background-color: #10B981;
            color: white;
            border: none;
            padding: 8px 16px;
            border-radius: 6px;
            font-weight: bold;
        }
        QPushButton:hover { background-color: #059669; }
        QPushButton:pressed { background-color: #047857; }
        """

        browse_button_style = """
        QPushButton {
            background-color: #374151;
            color: white;
            border: none;
            padding: 8px 16px;
            border-radius: 6px;
            font-weight: bold;
        }
        QPushButton:hover { background-color: #4B5563; }
        QPushButton:pressed { background-color: #1F2937; }
        """

        cancel_button_style = """
        QPushButton {
            background-color: transparent;
            color: #6B7280;
            border: 1px solid #D1D5DB;
            padding: 8px 16px;
            border-radius: 6px;
            font-weight: bold;
        }
        QPushButton:hover {
            background-color: #F3F4F6;
            color: #4B5563;
            border-color: #9CA3AF;
        }
        QPushButton:pressed {
            background-color: #E5E7EB;
            color: #1F2937;
        }
        """

        restore_button_style = """
        QPushButton {
            background-color: transparent;
            color: #EF4444;
            border: 1px solid #EF4444;
            padding: 8px 16px;
            border-radius: 6px;
            font-weight: bold;
        }
        QPushButton:hover {
            background-color: #FEF2F2;
            border-color: #DC2626;
            color: #DC2626;
        }
        QPushButton:pressed {
            background-color: #FEE2E2;
            color: #B91C1C;
            border-color: #B91C1C;
        }
        """

        main_layout = QVBoxLayout(self)

        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll_area.setFrameShape(QFrame.NoFrame)

        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        content_layout.setContentsMargins(5, 5, 5, 5)
        content_layout.setSpacing(10)

        self.tab_widget = QTabWidget()

        general_tab = QWidget()
        general_layout = QVBoxLayout(general_tab)
        general_layout.setContentsMargins(5, 5, 5, 5)
        general_layout.setSpacing(15)

        interface_group = QGroupBox(self.translate_text("Interface"))
        interface_layout = QFormLayout(interface_group)
        interface_layout.setVerticalSpacing(10)

        self.auto_open_checkbox = AnimatedCheckBox(self.translate_text("Ouvrir automatiquement le dernier projet au démarrage"))
        self.auto_open_checkbox.setChecked(self.config.get("auto_open_last_project", False))

        self.notifications_checkbox = AnimatedCheckBox(self.translate_text("Activer les notifications"))
        self.notifications_checkbox.setChecked(self.config.get("enable_notifications", True))

        self.system_notifications_checkbox = AnimatedCheckBox(self.translate_text("Activer les notifications système"))
        self.system_notifications_checkbox.setChecked(self.config.get("enable_system_notifications", True))

        self.show_previews_checkbox = AnimatedCheckBox(self.translate_text("Afficher les aperçus des fichiers"))
        self.show_previews_checkbox.setChecked(self.config.get("show_file_previews", True))

        self.show_dashboard_checkbox = AnimatedCheckBox(self.translate_text("Afficher le tableau de bord au démarrage"))
        self.show_dashboard_checkbox.setChecked(self.config.get("show_dashboard_on_startup", True))

        self.separate_image_pdfs_checkbox = AnimatedCheckBox(self.translate_text("Créer un PDF séparé par image (au lieu de fusionner)"))
        self.separate_image_pdfs_checkbox.setChecked(self.config.get("separate_image_pdfs", False))

        interface_layout.addRow(self.auto_open_checkbox)
        interface_layout.addRow(self.notifications_checkbox)
        interface_layout.addRow(self.system_notifications_checkbox)
        interface_layout.addRow(self.show_previews_checkbox)
        interface_layout.addRow(self.show_dashboard_checkbox)
        interface_layout.addRow(self.separate_image_pdfs_checkbox)

        self.use_system_theme_checkbox = AnimatedCheckBox(self.translate_text("Utiliser le thème système"))
        self.use_system_theme_checkbox.setChecked(self.config.get("use_system_theme", True))
        interface_layout.addRow(self.use_system_theme_checkbox)

        conversion_group = QGroupBox(self.translate_text("Conversion"))
        conversion_layout = QFormLayout(conversion_group)
        conversion_layout.setVerticalSpacing(12)
        conversion_layout.setHorizontalSpacing(15)

        self.quality_combo = QComboBox()
        self.quality_combo.addItems([
            self.translate_text("Haute qualité"),
            self.translate_text("Qualité standard"),
            self.translate_text("Compressé")
        ])
        quality_map = {"high": 0, "standard": 1, "compressed": 2}
        current_quality = self.config.get("conversion_quality", "standard")
        self.quality_combo.setCurrentIndex(quality_map.get(current_quality, 1))

        self.compression_level_combo = QComboBox()
        self.compression_level_combo.addItems([
            self.translate_text("Normal"),
            self.translate_text("Haute compression"),
            self.translate_text("Compression maximale")
        ])
        compression_map = {"normal": 0, "high": 1, "maximum": 2}
        current_compression = self.config.get("compression_level", "normal")
        self.compression_level_combo.setCurrentIndex(compression_map.get(current_compression, 0))

        self.pdf_to_word_mode_combo = QComboBox()
        self.pdf_to_word_mode_combo.addItems([
            self.translate_text("Conserver les images et la mise en page"),
            self.translate_text("Texte brut uniquement"),
            self.translate_text("Texte complet (texte + texte des images)")
        ])
        mode_map = {"with_images": 0, "text_only": 1, "text_with_image_text": 2}
        current_mode = self.config.get("pdf_to_word_mode", "with_images")
        self.pdf_to_word_mode_combo.setCurrentIndex(mode_map.get(current_mode, 0))

        self.default_output_input = QLineEdit()
        self.default_output_input.setText(self.config.get("default_output_folder", ""))
        self.browse_btn = QPushButton(self.translate_text("Parcourir"))
        self.browse_btn.setStyleSheet(browse_button_style)
        self.browse_btn.clicked.connect(self.browse_output_folder)

        output_layout = QHBoxLayout()
        output_layout.addWidget(self.default_output_input, 3)
        output_layout.addWidget(self.browse_btn, 1)
        output_layout.setSpacing(10)

        self.auto_clean_checkbox = AnimatedCheckBox(self.translate_text("Nettoyer automatiquement les fichiers temporaires"))
        self.auto_clean_checkbox.setChecked(self.config.get("auto_clean_temp_files", True))

        self.backup_checkbox = AnimatedCheckBox(self.translate_text("Créer une sauvegarde avant conversion"))
        self.backup_checkbox.setChecked(self.config.get("backup_before_conversion", False))

        self.keep_history_days_spin = QSpinBox()
        self.keep_history_days_spin.setRange(1, 3650)
        self.keep_history_days_spin.setValue(self.config.get("keep_history_days", 365))

        self.auto_save_templates_check = AnimatedCheckBox(self.translate_text("Sauvegarder automatiquement les configurations fréquentes"))
        self.auto_save_templates_check.setChecked(self.config.get("auto_save_templates", True))

        conversion_layout.addRow(self.translate_text("Qualité par défaut:"), self.quality_combo)
        conversion_layout.addRow(self.translate_text("Niveau compression:"), self.compression_level_combo)
        conversion_layout.addRow(self.translate_text("Mode PDF→Word:"), self.pdf_to_word_mode_combo)
        conversion_layout.addRow(self.translate_text("Dossier de sortie par défaut:"), output_layout)
        conversion_layout.addRow(self.translate_text("Conserver l'historique (jours):"), self.keep_history_days_spin)
        conversion_layout.addRow(self.auto_clean_checkbox)
        conversion_layout.addRow(self.backup_checkbox)
        conversion_layout.addRow(self.auto_save_templates_check)

        conversion_group.setMaximumHeight(400)

        general_layout.addWidget(interface_group)
        general_layout.addWidget(conversion_group)
        general_layout.addStretch()

        privacy_tab = QWidget()
        privacy_layout = QVBoxLayout(privacy_tab)
        privacy_layout.setContentsMargins(5, 5, 5, 5)
        privacy_layout.setSpacing(15)

        privacy_group = QGroupBox(self.translate_text("Confidentialité et Conditions"))
        privacy_group_layout = QVBoxLayout(privacy_group)
        privacy_group_layout.setSpacing(10)

        self.view_terms_btn = QPushButton("📄 " + self.translate_text("Voir les Conditions d'utilisation"))
        self.view_terms_btn.setMinimumHeight(33)
        self.view_privacy_btn = QPushButton("🔒 " + self.translate_text("Voir la Politique de confidentialité"))
        self.view_privacy_btn.setMinimumHeight(33)

        privacy_group_layout.addWidget(self.view_terms_btn)
        privacy_group_layout.addWidget(self.view_privacy_btn)

        privacy_layout.addWidget(privacy_group)
        privacy_layout.addStretch()

        self.tab_widget.addTab(general_tab,  self.translate_text("Général"))
        self.tab_widget.addTab(privacy_tab,  self.translate_text("Confidentialité"))
        self.tab_widget.addTab(self._build_language_tab(), self.translate_text("Langue"))

        content_layout.addWidget(self.tab_widget)

        scroll_area.setWidget(content_widget)
        main_layout.addWidget(scroll_area)

        ok_btn = QPushButton(self.translate_text("OK"))
        ok_btn.setStyleSheet(ok_button_style)
        cancel_btn = QPushButton(self.translate_text("Annuler"))
        cancel_btn.setStyleSheet(cancel_button_style)
        restore_btn = QPushButton(self.translate_text("Restaurer les paramètres par défaut"))
        restore_btn.setStyleSheet(restore_button_style)

        buttons = QDialogButtonBox()
        buttons.addButton(ok_btn, QDialogButtonBox.AcceptRole)
        buttons.addButton(cancel_btn, QDialogButtonBox.RejectRole)
        buttons.addButton(restore_btn, QDialogButtonBox.ResetRole)

        ok_btn.clicked.connect(self.accept)
        cancel_btn.clicked.connect(self.reject)
        restore_btn.clicked.connect(self.restore_defaults)

        main_layout.addWidget(buttons)

        self.view_terms_btn.clicked.connect(self.show_terms)
        self.view_privacy_btn.clicked.connect(self.show_privacy)

        self.apply_scrollbar_style()

    def apply_scrollbar_style(self):
        if hasattr(self.parent(), 'dark_mode') and self.parent().dark_mode:
            scrollbar_style = """
            QScrollBar:vertical {
                background-color: #4d5564; width: 10px; margin: 0px; border: none;
            }
            QScrollBar::handle:vertical {
                background-color: #6c757d; border-radius: 5px; min-height: 30px;
            }
            QScrollBar::handle:vertical:hover { background-color: #868e96; }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0px; }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background-color: #4d5564; }
            QScrollBar:horizontal {
                background-color: #4d5564; height: 10px; margin: 0px; border: none;
            }
            QScrollBar::handle:horizontal {
                background-color: #6c757d; border-radius: 5px; min-width: 30px;
            }
            QScrollBar::handle:horizontal:hover { background-color: #868e96; }
            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal { background-color: #4d5564; }
            """
        else:
            scrollbar_style = """
            QScrollBar:vertical {
                background-color: #dee2e6; width: 10px; margin: 0px; border: none;
            }
            QScrollBar::handle:vertical {
                background-color: #adb5bd; border-radius: 5px; min-height: 30px;
            }
            QScrollBar::handle:vertical:hover { background-color: #868e96; }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0px; }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background-color: #dee2e6; }
            QScrollBar:horizontal {
                background-color: #dee2e6; height: 10px; margin: 0px; border: none;
            }
            QScrollBar::handle:horizontal {
                background-color: #adb5bd; border-radius: 5px; min-width: 30px;
            }
            QScrollBar::handle:horizontal:hover { background-color: #868e96; }
            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal { background-color: #dee2e6; }
            """
        self.setStyleSheet(self.styleSheet() + scrollbar_style)

    def translate_text(self, text):
        return self._tm.translate_text(text)

    # Language tab
    def _lang_is_dark(self) -> bool:
        """Return True if the parent app is currently in dark mode."""
        p = self.parent()
        return bool(getattr(p, "dark_mode", False))

    def _build_language_tab(self) -> QWidget:
        """Build the Language settings tab with installed-language list + import."""
        dark = self._lang_is_dark()

        # Palette
        # Backgrounds
        tab_bg        = "#0d1117" if dark else "#f8f9fa"
        group_bg      = "#161b22" if dark else "#ffffff"
        scroll_bg     = "#161b22" if dark else "#ffffff"
        # Text
        text_primary  = "#e6edf3" if dark else "#1c2526"
        text_muted    = "#8b949e" if dark else "#6b7280"
        # Borders
        group_border  = "#30363d" if dark else "#dee2e6"
        # Active indicator
        active_green  = "#3fb950" if dark else "#10B981"

        tab = QWidget()
        tab.setStyleSheet(f"background-color: {tab_bg};")
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(8, 10, 8, 10)
        layout.setSpacing(12)

        # Active language indicator
        active_row = QHBoxLayout()
        active_lbl_title = QLabel(self.translate_text("Langue active:"))
        active_lbl_title.setStyleSheet(f"color: {text_primary}; font-size: 12px;")
        active_row.addWidget(active_lbl_title)
        self._lang_active_lbl = QLabel()
        self._lang_active_lbl.setStyleSheet(
            f"font-weight: bold; font-size: 13px; color: {active_green};"
        )
        active_row.addWidget(self._lang_active_lbl)
        active_row.addStretch()
        layout.addLayout(active_row)

        # Language list
        list_group = QGroupBox(self.translate_text("Langues installées"))
        list_group.setStyleSheet(f"""
            QGroupBox {{
                color: {text_muted};
                background-color: {group_bg};
                border: 1px solid {group_border};
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
                font-size: 11px;
                font-weight: 600;
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 6px;
                color: {text_muted};
            }}
        """)
        list_layout = QVBoxLayout(list_group)
        list_layout.setSpacing(6)
        list_layout.setContentsMargins(8, 8, 8, 8)

        self._lang_list = QScrollArea()
        self._lang_list.setWidgetResizable(True)
        self._lang_list.setFrameShape(QFrame.NoFrame)
        self._lang_list.setMinimumHeight(185)
        self._lang_list.setStyleSheet(f"""
            QScrollArea {{ background: {scroll_bg}; border: none; }}
            QScrollBar:vertical {{
                background: {"#21262d" if dark else "#e9ecef"};
                width: 6px; border-radius: 3px; margin: 0;
            }}
            QScrollBar::handle:vertical {{
                background: {"#484f58" if dark else "#adb5bd"};
                border-radius: 3px; min-height: 24px;
            }}
            QScrollBar::handle:vertical:hover {{
                background: {"#6e7681" if dark else "#6c757d"};
            }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height: 0; }}
        """)
        self._lang_inner = QWidget()
        self._lang_inner.setStyleSheet(f"background: {scroll_bg};")
        self._lang_inner_layout = QVBoxLayout(self._lang_inner)
        self._lang_inner_layout.setSpacing(6)
        self._lang_inner_layout.setContentsMargins(4, 4, 4, 4)
        self._lang_list.setWidget(self._lang_inner)
        list_layout.addWidget(self._lang_list)
        layout.addWidget(list_group)

        # Import button
        import_btn = QPushButton("📥 " + self.translate_text("Importer un fichier .lang"))
        import_btn.setMinimumHeight(36)
        import_btn.setStyleSheet("""
            QPushButton {
                background-color: #6366F1; color: white;
                border: none; padding: 8px 16px;
                border-radius: 6px; font-weight: bold; font-size: 12px;
            }
            QPushButton:hover   { background-color: #4F46E5; }
            QPushButton:pressed { background-color: #4338CA; }
        """)
        import_btn.clicked.connect(self._import_lang_file)
        layout.addWidget(import_btn)

        # Hint
        hint = QLabel("ℹ️ " + self.translate_text("lang_restart_hint"))
        hint.setWordWrap(True)
        hint.setStyleSheet(
            f"color: {text_muted}; font-size: 10px; font-style: italic;"
        )
        layout.addWidget(hint)
        layout.addStretch()

        self._refresh_lang_list()
        return tab

    def _refresh_lang_list(self) -> None:
        """Re-populate the language cards list."""
        from translations import TranslationManager

        while self._lang_inner_layout.count():
            item = self._lang_inner_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        parent_app = self.parent()
        if hasattr(parent_app, "translation_manager"):
            tm = parent_app.translation_manager
        else:
            tm = TranslationManager()

        current = getattr(parent_app, "current_language", "fr")
        langs   = tm.get_available_languages()
        active_name = next((l["name"] for l in langs if l["code"] == current), current)
        self._lang_active_lbl.setText(active_name)

        for lang in langs:
            card = self._make_lang_card(lang, current)
            self._lang_inner_layout.addWidget(card)

        self._lang_inner_layout.addStretch()

    def _make_lang_card(self, lang: dict, current_code: str) -> QFrame:
        """Build one language card — fully theme-aware."""
        dark = self._lang_is_dark()

        is_active  = lang["code"] == current_code
        is_builtin = lang.get("builtin", False)

        # Card palette
        if dark:
            if is_active:
                card_bg     = "#0f2a1e"   # deep green tint
                border_col  = "#3fb950"   # bright green
                name_col    = "#e6edf3"   # white-ish
                meta_col    = "#8b949e"   # muted grey
            else:
                card_bg     = "#1c2333"   # raised surface
                border_col  = "#30363d"   # subtle border
                name_col    = "#c9d1d9"   # light grey
                meta_col    = "#8b949e"   # muted grey
        else:
            if is_active:
                card_bg     = "#f0fdf4"   # mint tint
                border_col  = "#10B981"   # green
                name_col    = "#064e3b"   # dark green text
                meta_col    = "#6b7280"   # grey
            else:
                card_bg     = "#ffffff"
                border_col  = "#dee2e6"
                name_col    = "#1c2526"   # near-black
                meta_col    = "#6b7280"

        card = QFrame()
        card.setFrameShape(QFrame.StyledPanel)
        card.setStyleSheet(f"""
            QFrame {{
                border: 2px solid {border_col};
                border-radius: 8px;
                background-color: {card_bg};
                padding: 2px;
            }}
        """)

        row = QHBoxLayout(card)
        row.setContentsMargins(10, 8, 10, 8)
        row.setSpacing(10)

        # Badge
        if is_active:
            badge_text  = self.translate_text("lang_active_badge")
            badge_bg    = "#3fb950" if dark else "#10B981"
            badge_fg    = "#ffffff"
        elif is_builtin:
            badge_text  = self.translate_text("lang_builtin_badge")
            badge_bg    = "#388bfd" if dark else "#6366F1"
            badge_fg    = "#ffffff"
        else:
            badge_text  = self.translate_text("lang_external_badge")
            badge_bg    = "#f0a030" if dark else "#f59e0b"
            badge_fg    = "#ffffff" if dark else "#1c2526"

        badge = QLabel(badge_text)
        badge.setStyleSheet(f"""
            background-color: {badge_bg};
            color: {badge_fg};
            padding: 2px 8px;
            border-radius: 10px;
            font-size: 10px;
            font-weight: bold;
            border: none;
        """)
        row.addWidget(badge)

        # Name + meta
        info_col = QVBoxLayout()
        info_col.setSpacing(2)

        name_lbl = QLabel(lang["name"])
        name_lbl.setStyleSheet(
            f"font-weight: bold; font-size: 13px; color: {name_col}; border: none;"
        )
        info_col.addWidget(name_lbl)

        meta_parts = []
        if lang.get("author"):
            meta_parts.append(
                f"{self.translate_text('lang_author_label')} {lang['author']}"
            )
        if lang.get("version"):
            meta_parts.append(
                f"{self.translate_text('lang_version_label')} {lang['version']}"
            )
        desc = lang.get("description") or self.translate_text("lang_no_desc")
        meta_parts.append(desc)

        meta_lbl = QLabel("  •  ".join(meta_parts))
        meta_lbl.setStyleSheet(
            f"color: {meta_col}; font-size: 10px; border: none;"
        )
        meta_lbl.setWordWrap(True)
        info_col.addWidget(meta_lbl)
        row.addLayout(info_col, 1)

        # Apply button
        if not is_active:
            if dark:
                apply_ss = """
                    QPushButton {
                        background-color: #1f3a6e; color: #79c0ff;
                        border: 1px solid #388bfd; padding: 5px 10px;
                        border-radius: 5px; font-size: 11px; font-weight: bold;
                    }
                    QPushButton:hover { background-color: #388bfd; color: #ffffff; }
                    QPushButton:pressed { background-color: #1158c7; color: #ffffff; }
                """
            else:
                apply_ss = """
                    QPushButton {
                        background-color: #eff6ff; color: #1d4ed8;
                        border: 1px solid #93c5fd; padding: 5px 10px;
                        border-radius: 5px; font-size: 11px; font-weight: bold;
                    }
                    QPushButton:hover { background-color: #1d4ed8; color: #ffffff; }
                    QPushButton:pressed { background-color: #1e40af; color: #ffffff; }
                """
            apply_btn = QPushButton(self.translate_text("Appliquer"))
            apply_btn.setFixedWidth(82)
            apply_btn.setStyleSheet(apply_ss)
            code = lang["code"]
            apply_btn.clicked.connect(lambda _, c=code: self._apply_language(c))
            row.addWidget(apply_btn)

        # Remove button (external only)
        if not is_builtin:
            if dark:
                remove_ss = """
                    QPushButton {
                        background-color: #3d1a1a; color: #f85149;
                        border: 1px solid #6e1a1a; padding: 4px;
                        border-radius: 5px; font-size: 13px;
                    }
                    QPushButton:hover { background-color: #f85149; color: #ffffff; }
                """
            else:
                remove_ss = """
                    QPushButton {
                        background-color: #fef2f2; color: #ef4444;
                        border: 1px solid #fca5a5; padding: 4px;
                        border-radius: 5px; font-size: 13px;
                    }
                    QPushButton:hover { background-color: #ef4444; color: #ffffff; }
                """
            remove_btn = QPushButton("🗑")
            remove_btn.setFixedWidth(32)
            remove_btn.setToolTip(self.translate_text("Supprimer"))
            remove_btn.setStyleSheet(remove_ss)
            code = lang["code"]
            name = lang["name"]
            remove_btn.clicked.connect(
                lambda _, c=code, n=name: self._remove_language(c, n)
            )
            row.addWidget(remove_btn)

        return card

    def _import_lang_file(self) -> None:
        """Open file dialog to import a .lang file."""
        filepath, _ = QFileDialog.getOpenFileName(
            self,
            self.translate_text("Choisir un fichier .lang"),
            "",
            self.translate_text("Fichiers de langue (*.lang)"),
        )
        if not filepath:
            return

        parent_app = self.parent()
        if hasattr(parent_app, "translation_manager"):
            tm = parent_app.translation_manager
        else:
            from translations import TranslationManager
            tm = TranslationManager()

        ok, result = tm.load_lang_file(filepath)
        if ok:
            QMessageBox.information(
                self,
                self.translate_text("Langue importée"),
                self.translate_text("lang_import_ok").format(name=result),
            )
            self._refresh_lang_list()
        else:
            QMessageBox.warning(
                self,
                self.translate_text("Erreur"),
                self.translate_text("lang_import_err").format(error=result),
            )

    def _apply_language(self, code: str) -> None:
        """Switch the application language and refresh the card list."""
        parent_app = self.parent()
        if hasattr(parent_app, "translation_manager"):
            parent_app.translation_manager.set_language(code)
        if hasattr(parent_app, "current_language"):
            parent_app.current_language = code
        if hasattr(parent_app, "config"):
            parent_app.config["language"] = code
        if hasattr(parent_app, "config_manager"):
            parent_app.config_manager.save_config(parent_app.config)
        if hasattr(parent_app, "update_texts"):
            parent_app.update_texts()
        # Update language toolbar button text if present
        if hasattr(parent_app, "language_action"):
            if code == "fr":
                parent_app.language_action.setText("🇬🇧 English")
            elif code == "en":
                parent_app.language_action.setText("🇫🇷 Français")
            else:
                langs = parent_app.translation_manager.get_available_languages()
                name  = next((l["name"] for l in langs if l["code"] == code), code)
                parent_app.language_action.setText(f"🌐 {name}")
        self._refresh_lang_list()

    def _remove_language(self, code: str, name: str) -> None:
        """Confirm then remove an external language."""
        reply = QMessageBox.question(
            self,
            self.translate_text("Confirmation"),
            self.translate_text("lang_remove_confirm").format(name=name),
            QMessageBox.Yes | QMessageBox.No,
        )
        if reply != QMessageBox.Yes:
            return

        parent_app = self.parent()
        if hasattr(parent_app, "translation_manager"):
            tm = parent_app.translation_manager
        else:
            from translations import TranslationManager
            tm = TranslationManager()

        ok, error = tm.remove_lang_file(code)
        if ok:
            QMessageBox.information(
                self,
                self.translate_text("Langue supprimée"),
                self.translate_text("lang_remove_ok").format(code=code),
            )
            self._refresh_lang_list()
        else:
            QMessageBox.warning(
                self,
                self.translate_text("Erreur"),
                self.translate_text("lang_remove_err").format(error=error),
            )

    def browse_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, self.translate_text("Sélectionner le dossier de sortie par défaut"))
        if folder:
            self.default_output_input.setText(folder)

    def show_terms(self):
        dark_mode = self.config.get("dark_mode", False)
        dialog = TermsAndPrivacyDialog(self, self.language, dark_mode=dark_mode)
        dialog.tab_widget.setCurrentIndex(0)
        dialog.from_settings = True
        result = dialog.exec()
        if dialog.closed_by_cross:
            return
        self.process_terms_result(result, dialog)

    def show_privacy(self):
        dark_mode = self.config.get("dark_mode", False)
        dialog = TermsAndPrivacyDialog(self, self.language, dark_mode=dark_mode)
        dialog.tab_widget.setCurrentIndex(1)
        dialog.from_settings = True
        result = dialog.exec()
        if dialog.closed_by_cross:
            return
        self.process_terms_result(result, dialog)

    def process_terms_result(self, result, dialog):
        parent_app = self.parent()
        if not hasattr(parent_app, 'config_manager'):
            return

        config = parent_app.config_manager.load_config()

        if result == QDialog.Accepted:
            config["accepted_terms"] = True
            config["accepted_privacy"] = True

            if config.get("terms_acceptance_timestamp") is not None:
                config["terms_reacceptance_timestamp"] = datetime.now().isoformat()
                print(f"[TERMS DEBUG] ✅ Re-acceptance detected - terms_reacceptance_timestamp added: {config['terms_reacceptance_timestamp']}")
            else:
                config["terms_acceptance_timestamp"] = datetime.now().isoformat()
                print(f"[TERMS DEBUG] ℹ️ First acceptance - terms_acceptance_timestamp set: {config['terms_acceptance_timestamp']}")

            parent_app.config_manager.save_config(config)
            parent_app.config.update(config)
            parent_app.terms_accepted = True

            QMessageBox.information(
                self,
                self.translate_text("Succès"),
                self.translate_text("Conditions et politique acceptées avec succès.")
            )

            self.accept()
        else:
            config["accepted_terms"] = False
            config["accepted_privacy"] = False
            config["terms_rejection_timestamp"] = datetime.now().isoformat()

            parent_app.config_manager.save_config(config)
            parent_app.config.update(config)

            QMessageBox.warning(
                self,
                self.translate_text("Attention"),
                self.translate_text("Vous avez refusé les conditions. L'application va se fermer.")
            )

            QTimer.singleShot(1500, QCoreApplication.quit)
            return

    def restore_defaults(self):
        self.auto_open_checkbox.setChecked(False)
        self.notifications_checkbox.setChecked(True)
        self.system_notifications_checkbox.setChecked(True)
        self.show_previews_checkbox.setChecked(True)
        self.show_dashboard_checkbox.setChecked(True)
        self.quality_combo.setCurrentIndex(1)
        self.compression_level_combo.setCurrentIndex(0)
        self.pdf_to_word_mode_combo.setCurrentIndex(0)
        self.default_output_input.clear()
        self.keep_history_days_spin.setValue(365)
        self.auto_clean_checkbox.setChecked(True)
        self.backup_checkbox.setChecked(False)
        self.auto_save_templates_check.setChecked(True)
        self.separate_image_pdfs_checkbox.setChecked(False)
        self.use_system_theme_checkbox.setChecked(True)

    def get_settings(self):
        quality_map = {0: "high", 1: "standard", 2: "compressed"}
        compression_map = {0: "normal", 1: "high", 2: "maximum"}
        mode_map = {0: "with_images", 1: "text_only", 2: "text_with_image_text"}
        return {
            "auto_open_last_project": self.auto_open_checkbox.isChecked(),
            "enable_notifications": self.notifications_checkbox.isChecked(),
            "enable_system_notifications": self.system_notifications_checkbox.isChecked(),
            "show_file_previews": self.show_previews_checkbox.isChecked(),
            "show_dashboard_on_startup": self.show_dashboard_checkbox.isChecked(),
            "conversion_quality": quality_map[self.quality_combo.currentIndex()],
            "compression_level": compression_map[self.compression_level_combo.currentIndex()],
            "pdf_to_word_mode": mode_map[self.pdf_to_word_mode_combo.currentIndex()],
            "default_output_folder": self.default_output_input.text(),
            "keep_history_days": self.keep_history_days_spin.value(),
            "auto_clean_temp_files": self.auto_clean_checkbox.isChecked(),
            "backup_before_conversion": self.backup_checkbox.isChecked(),
            "auto_save_templates": self.auto_save_templates_check.isChecked(),
            "separate_image_pdfs": self.separate_image_pdfs_checkbox.isChecked(),
            "use_system_theme": self.use_system_theme_checkbox.isChecked()
        }

class ConversionOptionsDialog(QDialog):
    """
    Dialog box allowing the user to choose the type of conversion
    to launch via a keyboard shortcut or a generic button.
    Premium redesign: card-based grid, large icons, dark/light theme-aware.
    """
    conversion_chosen = Signal(str)

    # (label_key, method_name, icon_file, accent_dark, accent_light)
    _CARDS = [
        ("PDF → Word/Docx",       "launch_pdf_to_word_conversion",  "pdf_word.png",        "#4dabf7", "#1971c2"),
        ("Word/Docx → PDF",       "launch_word_to_pdf_conversion",  "word_pdf.png",        "#69db7c", "#2f9e44"),
        ("Images → PDF",          "launch_image_to_pdf_conversion", "image_pdf.png",       "#ffa94d", "#e67700"),
        ("Fusionner PDF",         "launch_merge_pdf",               "merge_pdf.png",       "#f783ac", "#c2255c"),
        ("Fusionner Word",        "launch_merge_word",              "merge_word.png",      "#da77f2", "#9c36b5"),
        ("Optimiser les fichiers","launch_office_optimization",     "compress_pdf.png",    "#63e6be", "#0ca678"),
        ("Plus de conversions",   "launch_more_conversions",        "more_conversions.png","#74c0fc", "#1971c2"),
    ]

    def __init__(self, parent=None):
        super().__init__(parent)
        self._dark = getattr(parent, "dark_mode", True)
        self.setWindowTitle(parent.translate_text("Lancer la Conversion"))
        self.setFixedWidth(560)

        flags = (Qt.Window | Qt.WindowTitleHint | Qt.WindowSystemMenuHint | Qt.WindowCloseButtonHint)
        self.setWindowFlags(flags | Qt.CustomizeWindowHint)

        # Global color palette for the dialog
        if self._dark:
            dlg_bg     = "#0d1117"
            dlg_border = "#30363d"
            title_col  = "#e6edf3"
            sub_col    = "#8b949e"
            cancel_bg  = "#21262d"
            cancel_fg  = "#c9d1d9"
            cancel_brd = "#30363d"
            cancel_hov = "#30363d"
        else:
            dlg_bg     = "#f5f5f5"
            dlg_border = "#d0d7de"
            title_col  = "#1f2328"
            sub_col    = "#656d76"
            cancel_bg  = "#f6f8fa"
            cancel_fg  = "#24292f"
            cancel_brd = "#d0d7de"
            cancel_hov = "#e9ecef"

        self.setStyleSheet(f"""
            QDialog {{
                background-color: {dlg_bg};
                border: 1px solid {dlg_border};
                border-radius: 12px;
            }}
        """)

        root = QVBoxLayout(self)
        root.setContentsMargins(24, 20, 24, 20)
        root.setSpacing(0)

        # Header
        title = QLabel(parent.translate_text("Lancer la Conversion"))
        title.setStyleSheet(
            f"font-size: 17px; font-weight: 700; color: {title_col}; background: transparent;"
        )
        root.addWidget(title)

        sub = QLabel(parent.translate_text("Sélectionnez la catégorie de conversion à lancer :"))
        sub.setStyleSheet(
            f"font-size: 12px; color: {sub_col}; background: transparent; margin-bottom: 14px;"
        )
        root.addWidget(sub)
        root.addSpacing(10)

        # Cards grid
        grid = QGridLayout()
        grid.setHorizontalSpacing(10)
        grid.setVerticalSpacing(10)

        row, col = 0, 0
        for label_key, method_name, icon_file, accent_dark, accent_light in self._CARDS:
            label   = parent.translate_text(label_key)
            accent  = accent_dark if self._dark else accent_light
            is_more = (method_name == "launch_more_conversions")
            card    = self._make_card(label, icon_file, method_name, accent, is_more)
            if is_more:
                grid.addWidget(card, row, 0, 1, 2)
                row += 1
                col  = 0
            else:
                grid.addWidget(card, row, col)
                col += 1
                if col > 1:
                    col = 0
                    row += 1

        root.addLayout(grid)
        root.addSpacing(16)

        # Cancel
        cancel = QPushButton(parent.translate_text("Annuler"))
        cancel.setFixedHeight(36)
        cancel.setStyleSheet(f"""
            QPushButton {{
                background-color: {cancel_bg};
                color: {cancel_fg};
                border: 1px solid {cancel_brd};
                border-radius: 8px;
                font-size: 13px;
                font-weight: 500;
            }}
            QPushButton:hover  {{ background-color: {cancel_hov}; }}
            QPushButton:pressed {{ opacity: 0.7; }}
        """)
        cancel.clicked.connect(self.reject)
        root.addWidget(cancel)

    # Card factory
    def _make_card(self, label: str, icon_file: str, method_name: str,
                   accent: str, full_width: bool) -> QPushButton:
        """Build one conversion card button."""
        if self._dark:
            card_bg  = "#161b22"
            card_brd = "#30363d"
            card_fg  = "#e6edf3"
            hov_bg   = "#1c2333"
        else:
            card_bg  = "#ffffff"
            card_brd = "#d0d7de"
            card_fg  = "#1f2328"
            hov_bg   = "#f0f6fc"

        btn = QPushButton()
        btn.setObjectName(method_name)
        btn.setCursor(Qt.PointingHandCursor)
        btn.setFixedHeight(48 if full_width else 72)

        text_align   = "center" if full_width else "left"
        padding_left = "0px"    if full_width else "14px"

        btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {card_bg};
                color: {card_fg};
                border: 1.5px solid {card_brd};
                border-radius: 10px;
                font-size: 13px;
                font-weight: 600;
                text-align: {text_align};
                padding-left: {padding_left};
            }}
            QPushButton:hover {{
                background-color: {hov_bg};
                border-color: {accent};
                color: {accent};
            }}
            QPushButton:pressed {{
                background-color: {hov_bg};
                border-color: {accent};
            }}
        """)

        # Icon
        icon_path = self.parent().get_resource_path(os.path.join("icons", icon_file))
        if os.path.exists(icon_path):
            btn.setIcon(QIcon(icon_path))
            btn.setIconSize(QSize(28, 28) if full_width else QSize(36, 36))

        btn.setText(("  " if os.path.exists(icon_path) else "") + label)

        # Connection
        if method_name == "launch_more_conversions":
            btn.clicked.connect(self._on_more_conversions)
        else:
            btn.clicked.connect(lambda _=False, m=method_name: (
                self.conversion_chosen.emit(m),
                self.accept()
            ))

        return btn

    def _on_more_conversions(self):
        from advanced_conversions import AdvancedConversionsDialog
        dialog = AdvancedConversionsDialog(self.parent(), self.parent().current_language)
        dialog.show()
        dialog.raise_()
        dialog.activateWindow()
        self.accept()

    # Compat alias
    def on_more_conversions_clicked(self):
        self._on_more_conversions()

    def create_option_button(self, text, icon_file, method_name):
        """Compatibility alias (old API)."""
        dark  = getattr(self.parent(), "dark_mode", True)
        acc_d = next((a for l, m, i, a, _ in self._CARDS if m == method_name), "#4dabf7")
        acc_l = next((b for l, m, i, _, b in self._CARDS if m == method_name), "#1971c2")
        return self._make_card(text, icon_file, method_name,
                               acc_d if dark else acc_l,
                               method_name == "launch_more_conversions")

from .word_to_pdf_dialog import WordToPdfOptionsDialog